"""학위논문 docx 통합 스크립트.

Markdown chapter 파일들을 학위논문 작성지침의 양식에 따라 통합 docx로 빌드.

양식 (학위논문지침 ⑷):
  - 큰제목: 16pt 진하게 (명조)
  - 중간제목: 13pt 진하게
  - 본문: 11pt
  - 각주: 9pt
  - 줄간격: 200%

구성:
  1. 겉표지 (별지 3호)
  2. 속표지 (별지 3호 동일)
  3. 제출문 (별지 4호)
  4. 논문 인준서 (별지 5호)
  5. 감사의 글
  6. 목차 (필요시 수동 갱신)
  7. 표 차례 / 그림 차례
  8. 영문 Abstract
  9. 국문 초록
  10. 본문 (1~8장)
  11. 참고문헌
  12. 부록

산출:
  thesis/draft/thesis_draft.docx

실행:
  PYTHONIOENCODING=utf-8 .venv/Scripts/python.exe -m thesis.draft.build_thesis_docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─────────────────────────────────────────────────────────────
# 상수
# ─────────────────────────────────────────────────────────────
DRAFT_DIR = Path(__file__).resolve().parent
KOREAN_FONT = "맑은 고딕"  # Malgun Gothic — Windows 기본, 학위논문지침 허용 범위
ENG_FONT = "Times New Roman"
LINE_SPACING = 1.6  # 160% (학위논문지침 최소값, 손지민 양식 유사)

# 정보
TITLE = "정형 데이터 특화 딥러닝(TabNet)과 거대언어모델(LLM) 기반 XAI-RAG를 활용한 설명 가능한 신용 평가 및 사용자 맞춤형 리포트 생성"
SUBTITLE = "어텐션-SHAP 융합 컨텍스트와 다중 평가 프레임워크 적용"
AUTHOR_KR = "오 현 택"
AUTHOR_EN = "Hyuntaek Oh"
ADVISOR_KR = "박 운 상"
ADVISOR_EN = "Unsang Park"
DEPT_KR = "데이터사이언스·인공지능 전공"
DEPT_EN = "Department of Data Science · AI"
SCHOOL_KR = "서강대학교 AI·SW대학원"
SCHOOL_EN = "Sogang University Graduate School of AI & SW"
SUBMIT_DATE_KR = "2026년 6월"
SUBMIT_DATE_FULL = "2026년 6월 일"


# ─────────────────────────────────────────────────────────────
# 유틸 — 폰트·간격 적용
# ─────────────────────────────────────────────────────────────
def set_run(run, font_kr=KOREAN_FONT, font_en=ENG_FONT, size_pt=11, bold=False):
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_kr)
    rFonts.set(qn("w:hAnsi"), font_en)
    rFonts.set(qn("w:ascii"), font_en)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_format(p, line_spacing=LINE_SPACING, alignment=None,
                          space_before=0, space_after=0, first_indent=0):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(first_indent)


def add_para(doc, text, size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
              indent=0, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size_pt=size, bold=bold)
    set_paragraph_format(p, alignment=alignment, space_before=space_before,
                         space_after=space_after, first_indent=indent)
    return p


def add_heading(doc, text, level=1):
    """level=1: 큰제목 16pt, level=2: 중간제목 13pt, level=3: 11pt 진하게.

    Word의 Heading 1/2/3 스타일을 적용하여 [참조 → 목차]로 자동 갱신 가능하게 한다.
    Heading 스타일을 적용한 후 한국어 폰트와 사이즈를 추가로 설정한다.
    """
    sizes = {1: 16, 2: 13, 3: 11}
    size = sizes.get(level, 11)
    style_name = f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    set_run(run, size_pt=size, bold=True)
    # Heading 스타일의 검은색 강제 (기본은 파란색일 수 있음)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=24, space_after=18)
    elif level == 2:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=18, space_after=12)
    else:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=12, space_after=6)
    return p


def add_image(doc, image_path: str, caption: str, width_cm: float = 14.0):
    """Markdown ![caption](path) → docx 이미지 + 캡션."""
    from pathlib import Path as _P
    p_img = _P(image_path)
    if not p_img.is_absolute():
        # build 스크립트 기준 상대 경로 → 프로젝트 루트 기준
        p_img = (Path(__file__).resolve().parent / image_path).resolve()
    if not p_img.exists():
        # 그림 없는 경우 placeholder 텍스트
        add_para(doc, f"[그림 누락: {image_path}]", size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_para(doc, caption, size=10, bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        return
    # 이미지 삽입 (단락 + 중앙 정렬)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(p_img), width=Cm(width_cm))
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=12, space_after=0, line_spacing=1.0)
    # 캡션 (10pt, 진하게, 중앙)
    add_para(doc, caption, size=10, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc_field(doc, instr=' TOC \\o "1-3" \\h \\z \\u ',
                    placeholder_text="(목차 — Word에서 마우스 우클릭 → '필드 업데이트' 또는 F9)"):
    """Word ToC 필드 삽입. 열 때 자동 갱신 가능."""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    set_run(run, size_pt=11)
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = placeholder_text
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(placeholder)
    r.append(fldChar_end)
    return p


# ─────────────────────────────────────────────────────────────
# 별지 양식 (3호 겉표지, 4호 제출문, 5호 인준서)
# ─────────────────────────────────────────────────────────────
def add_cover_page(doc):
    """별지 3호 — 겉표지."""
    # 5cm 위에서 제목
    for _ in range(5):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, TITLE, size=20, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
    add_para(doc, SUBTITLE, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

    for _ in range(8):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, SUBMIT_DATE_KR, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, SCHOOL_KR, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, DEPT_KR, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, AUTHOR_KR, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_page_break(doc)


def add_inner_cover(doc):
    """속표지 — 겉표지와 동일."""
    add_cover_page(doc)


def add_submission_page(doc):
    """별지 4호 — 제출문."""
    for _ in range(4):
        add_para(doc, "", size=11, space_after=0)
    add_para(doc, TITLE, size=20, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, SUBTITLE, size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_para(doc, f"지도교수    {ADVISOR_KR}", size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    for _ in range(6):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, "이 논문을 데이터사이언스·인공지능 석사 학위논문으로 제출함",
             size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, SUBMIT_DATE_FULL, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, SCHOOL_KR, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, DEPT_KR, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, AUTHOR_KR, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_page_break(doc)


def add_approval_page(doc):
    """별지 5호 — 논문 인준서."""
    for _ in range(4):
        add_para(doc, "", size=11, space_after=0)
    add_para(doc, "논 문 인 준 서", size=24, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    for _ in range(3):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, f"{AUTHOR_KR}의 데이터사이언스·인공지능 석사 학위논문을 인준함.",
             size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    for _ in range(4):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, SUBMIT_DATE_FULL, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    for _ in range(3):
        add_para(doc, "", size=11, space_after=0)

    add_para(doc, "주심                            (인)", size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "부심                            (인)", size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "부심                            (인)", size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_page_break(doc)


def add_acknowledgement(doc):
    """감사의 글 (간략 placeholder — 사용자가 직접 채울 수 있음)."""
    add_heading(doc, "감사의 글", level=1)
    text = (
        "본 논문이 완성되기까지 많은 분들의 도움이 있었습니다. "
        "먼저 끊임없는 지도와 격려를 보내주신 박운상 지도교수님께 진심으로 감사드립니다. "
        "또한 본 연구를 함께 검토해 주신 심사위원님들과 "
        "연구 과정에서 의견을 나눠준 동료 연구자들, "
        "그리고 항상 응원해 주신 가족에게 깊은 감사를 표합니다."
    )
    add_para(doc, text, size=11, space_after=12, indent=0.5)
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────
# Markdown → docx 변환
# ─────────────────────────────────────────────────────────────
def add_markdown_table(doc, lines):
    """간단 markdown table → docx table.

    lines: ['| col1 | col2 |', '|---|---|', '| v1 | v2 |', ...]
    """
    # 1: header, 2: separator, 3+: rows
    rows = []
    for i, line in enumerate(lines):
        if i == 1:
            continue  # separator
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_col = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_col)
    table.style = "Light Grid"
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            cell = table.rows[i].cells[j]
            # 기존 paragraph 비우고 새로
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            set_run(run, size_pt=10, bold=(i == 0))


def parse_chapter_markdown(doc, md_text: str):
    """Markdown chapter 파일을 docx로 변환.

    지원:
      # → 큰제목 (level 1)
      ## → 중간제목 (level 2)
      ### → 소제목 (level 3)
      | a | b | → 표
      ```...``` → 코드 블록 (작은 글씨)
      일반 단락
    """
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # 빈 줄
        if not line.strip():
            i += 1
            continue
        # 헤딩
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        # 코드 블록
        if line.startswith("```"):
            j = i + 1
            block = []
            while j < len(lines) and not lines[j].startswith("```"):
                block.append(lines[j])
                j += 1
            code = "\n".join(block)
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            set_paragraph_format(p, line_spacing=1.15, space_after=6)
            i = j + 1
            continue
        # Markdown 표
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            block = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith("|"):
                block.append(lines[j])
                j += 1
            add_markdown_table(doc, block)
            i = j
            continue
        # 이미지 ![caption](path)
        img_match = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if img_match:
            caption, img_path = img_match.group(1), img_match.group(2)
            add_image(doc, img_path, caption)
            i += 1
            continue
        # 인용 (>) — 수식 블록 ("> **수식 N-N**: ...") 은 박스로 처리
        if line.startswith("> "):
            quote_text = line[2:].strip()
            # **bold** 마크 제거
            quote_text = re.sub(r"\*\*(.+?)\*\*", r"\1", quote_text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(quote_text)
            set_run(run, size_pt=11, bold=False)
            run.italic = True
            set_paragraph_format(p, line_spacing=1.5,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=6, space_after=6,
                                 first_indent=0)
            i += 1
            continue
        # 일반 단락 (markdown 강조 제거)
        text = line
        # **bold** -> 그대로 (단순화), *italic* -> 그대로
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # IEEE 인용 [N]은 그대로
        # 리스트 아이템 - / *
        if line.startswith("- ") or line.startswith("* "):
            text = "• " + text[2:]
        elif re.match(r"^\d+\.\s", line):
            text = text  # 그대로
        add_para(doc, text, size=11, space_after=6, indent=0.5)
        i += 1


def add_chapter_from_md(doc, md_path: Path):
    text = md_path.read_text(encoding="utf-8")
    parse_chapter_markdown(doc, text)
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────
# 메인 빌드
# ─────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # 페이지 여백 (학위논문지침 외 기본)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    # 1~3. 표지 / 속표지 / 제출문 / 인준서
    add_cover_page(doc)
    add_inner_cover(doc)
    add_submission_page(doc)
    add_approval_page(doc)

    # 4. 감사의 글
    add_acknowledgement(doc)

    # 5. 목차 / 표 차례 / 그림 차례 — Word ToC 필드 사용 (열 때 자동 갱신)
    add_heading(doc, "목   차", level=1)
    add_toc_field(doc, instr=' TOC \\o "1-3" \\h \\z \\u ',
                   placeholder_text="(목차가 여기에 자동 생성됩니다. Word에서 마우스 우클릭 → 필드 업데이트 또는 F9)")
    add_page_break(doc)

    add_heading(doc, "표 차례", level=1)
    add_toc_field(doc, instr=' TOC \\h \\z \\c "표" ',
                   placeholder_text="(본문 캡션을 [참조 → 캡션 삽입]으로 등록 후 [F9]로 갱신)")
    add_page_break(doc)

    add_heading(doc, "그림 차례", level=1)
    add_toc_field(doc, instr=' TOC \\h \\z \\c "그림" ',
                   placeholder_text="(본문 캡션을 [참조 → 캡션 삽입]으로 등록 후 [F9]로 갱신)")
    add_page_break(doc)

    # 6. 영문 Abstract / 국문 초록
    add_chapter_from_md(doc, DRAFT_DIR / "abstract.md")

    # 7. 본문 1~8장
    for fn in [
        "chapter1_introduction.md",
        "chapter2_related_work.md",
        "chapter3_data.md",
        "chapter4_methodology.md",
        "chapter5_evaluation.md",
        "chapter6_results.md",
        "chapter7_discussion.md",
        "chapter8_conclusion.md",
    ]:
        add_chapter_from_md(doc, DRAFT_DIR / fn)

    # 8. 참고문헌
    add_chapter_from_md(doc, DRAFT_DIR / "references.md")

    # 9. 부록
    add_chapter_from_md(doc, DRAFT_DIR / "appendix.md")

    # 저장
    out = DRAFT_DIR / "thesis_draft.docx"
    doc.save(str(out))
    print(f"[OK] thesis docx 빌드 완료: {out}")
    print(f"     파일 크기: {out.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
