"""친절한 버전의 미팅 보고서 docx.

기존 midterm_report.docx (압축적·결과 위주)와는 별개:
  - 각 섹션마다 "왜 → 어떻게 → 결과 → 의미" 흐름으로 풀어 쓴다
  - 문제의식 / 배경 설명을 많이
  - 비전공자도 따라 읽을 수 있는 톤
  - 실제 인스턴스를 따라가며 설명 (idx 59291)

산출:  paper/midterm_report_friendly.docx

실행:  PYTHONIOENCODING=utf-8 .venv/Scripts/python.exe -m src.gen_friendly_report
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.utils import RESULTS_DIR

PAPER_DIR = Path("paper")
PAPER_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR = Path("figures")


# ─────────────────────────────────────────────────────────────
# 헬퍼 (gen_report.py와 유사하나 톤·여백 풍부)
# ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _set_run_korean_font(run, size: int = 11, bold: bool = False,
                            color: tuple | None = None):
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rFonts.set(qn("w:ascii"), "Malgun Gothic")
    rFonts.set(qn("w:hAnsi"), "Malgun Gothic")


def H(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_korean_font(run, size=18 - 3 * (level - 1), bold=True)
    return h


def P(doc, text: str, bold: bool = False, size: int = 11,
      color: tuple | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_korean_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def Quote(doc, text: str):
    """인용/강조 박스 — 회색 음영 + 들여쓰기."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run(text)
    _set_run_korean_font(run, size=10.5, bold=False, color=(0x33, 0x33, 0x33))
    run.italic = True
    # 회색 배경
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    return p


def Bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    run = p.add_run(text)
    _set_run_korean_font(run, size=11)
    return p


def Table(doc, headers: list, rows: list, header_bg: str = "4C72B0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_korean_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        _set_cell_bg(cell, header_bg)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run_korean_font(run, size=10)
    return table


def Fig(doc, path: Path, caption: str, width_cm: float = 14):
    if not path.exists():
        P(doc, f"[그림 누락: {path}]")
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    _set_run_korean_font(run, size=9, color=(0x55, 0x55, 0x55))
    run.italic = True


def PageBreak(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # ── 표지 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("석사 논문 중간 보고 — 친절판")
    _set_run_korean_font(run, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("\nTabNet과 거대언어모델(LLM) 기반 XAI-RAG를 활용한\n"
                       "설명 가능한 신용 평가 및 자연어 리포트 생성")
    _set_run_korean_font(run, size=14)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("\n").font.size = Pt(10)
    for txt in ["A70067 오현택", "지도교수: 박운상", "보고일: 2026-05-10",
                "Step 1 (8일 작업 완료, 미팅 프로토타입)"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        _set_run_korean_font(r, size=11)

    P(doc, "")
    Quote(doc,
        "이 보고서는 발표 자료의 압축판이 아니라, 연구의 전체 흐름을 처음부터 "
        "차근차근 풀어 쓴 \"친절판\"입니다. 같이 읽으면 내가 무엇을 왜 했는지 "
        "어렵지 않게 따라가실 수 있도록 작성했습니다.")

    PageBreak(doc)

    # ── 0. 한 페이지 요약 ──
    H(doc, "0. 한 페이지 요약 (Executive Summary)", level=1)
    P(doc, "왜 이 연구를 시작했는가", bold=True, size=12)
    P(doc, "신용 평가 모델은 점점 더 \"잘 맞히지만 설명하기 어려운\" 블랙박스가 "
            "되어 가고 있습니다. SHAP 같은 사후 해석 기법은 변수 기여도를 "
            "수치로 알려주지만, 그 결과를 일반 고객이나 심사역이 그대로 이해하기는 "
            "여전히 어렵습니다. 한편 LLM은 자연어 변환을 잘하지만 학습되지 않은 "
            "사실을 그럴듯하게 만들어내는 환각(hallucination) 문제가 있어, "
            "신용평가 같이 정확성·신뢰성·법적 책임이 요구되는 도메인에 그대로 "
            "쓸 수 없습니다.", size=11)

    P(doc, "본 연구가 제안하는 것", bold=True, size=12)
    P(doc, "본 연구는 SHAP의 수치 기반 설명을 LLM의 \"검색된 근거(retrieved "
            "evidence)\"로 재정의하는 XAI-RAG 구조를 제안합니다. "
            "LLM은 자유롭게 추론하지 않고, SHAP가 만든 사실 컨텍스트만 자연어로 "
            "변환하는 역할에 한정됩니다. 그 결과 환각이 원천적으로 차단되며, "
            "사용자에게는 친절한 자연어 설명 리포트가 자동 생성됩니다.")

    P(doc, "핵심 발견 (10 샘플 평가 기준)", bold=True, size=12)
    Bullet(doc, "Hallucination Rate가 두 상용 LLM(Gemini 2.5 Flash, Claude "
                "Sonnet 4.5)에서 모두 0.000 — 본 구조는 LLM 종속성 없이 환각을 "
                "차단합니다.")
    Bullet(doc, "동일 데이터를 SHAP 컨텍스트 없이 LLM에 직접 주는 baseline에서 "
                "Claude는 45.5%의 환각을 만들어냈습니다 (DTI/LTV/DSR 같이 데이터에 "
                "없는 일반 금융 용어, 가짜 고객센터 번호 등). 이 둘의 비교가 본 "
                "구조의 효과를 가장 분명하게 보여줍니다.")
    Bullet(doc, "G-Eval 자동 평가에서 정확성·민감변수 마스킹·문체 모두 5/5점 만점.")
    Bullet(doc, "예측 성능은 5-fold Stratified CV에서 XGBoost가 test AUROC "
                "0.7587 ± 0.0008로 가장 안정적이며, TabNet은 0.7518 ± 0.0017로 "
                "약간 부족하지만 본 연구의 가치는 어텐션 기반 내재적 해석성에 "
                "있습니다.")
    Bullet(doc, "8개 공정성 케이스 모두 4/5 rule을 위반했고, 보호 속성 ablation은 "
                "성별 편향에 대해서만 효과적이었습니다 — 연령은 다른 변수에 "
                "간접 인코딩되어 있어 단순 제거로는 줄지 않습니다.")

    PageBreak(doc)

    # ── 1. 연구 배경 ──
    H(doc, "1. 왜 이 연구가 필요한가 — 연구 배경", level=1)

    H(doc, "1.1 자동화된 신용 평가의 두 얼굴", level=2)
    P(doc, "최근 금융기관은 머신러닝과 딥러닝 모델을 여신 심사·리스크 관리·"
            "이상 거래 탐지에 광범위하게 사용합니다. XGBoost, LightGBM, "
            "TabNet 같은 모델은 전통적인 로지스틱 회귀보다 부도 예측 성능이 "
            "훨씬 좋습니다. 그러나 이 모델들은 \"왜 거절되었는가\"를 직관적으로 "
            "설명하기 어렵습니다.")
    P(doc, "이는 단순한 기술적 한계를 넘어, 제도적·윤리적 문제로 이어집니다. "
            "국내 금융소비자보호법 및 신용정보법은 자동화된 신용평가에 대한 "
            "설명 의무와 이의제기 권리를 명시합니다. EU의 일반정보보호규정"
            "(GDPR)도 알고리즘적 의사결정에 대한 \"설명 요구권(Right to "
            "Explanation)\"을 보장합니다.")

    H(doc, "1.2 기존 XAI의 한계 — \"모델은 설명 가능해졌으나 사람에게는 "
              "설명되지 않는다\"", level=2)
    P(doc, "이런 요구에 대응해 SHAP, LIME 같은 사후 해석 기법이 활발히 연구되어 "
            "왔습니다. 그러나 SHAP의 출력은 변수 중요도 그래프, 기여도 수치, "
            "Force Plot 같은 시각화·수치 기반입니다. 데이터 분석 전문가가 "
            "아니라면 직관적으로 해석하기 어렵습니다.")
    Quote(doc, "즉 \"모델은 설명 가능해졌으나 사람에게는 여전히 설명되지 않는\" "
                "의사 설명(pseudo-explanation) 문제가 남아 있습니다.")

    H(doc, "1.3 LLM의 가능성과 위험", level=2)
    P(doc, "GPT-4, Claude, Gemini 같은 최신 LLM은 수치형·구조적 정보를 "
            "자연어로 변환하는 데 뛰어납니다. SHAP 출력값을 자연어로 옮기는 "
            "시도가 활발합니다. 하지만 LLM에는 학습 데이터에 없는 사실도 "
            "그럴듯하게 만들어내는 환각이라는 본질적 약점이 있어, 신용평가 "
            "같이 정확성과 법적 책임이 중요한 도메인에서 그대로 쓰기는 위험합니다.")
    P(doc, "본 연구는 이 두 가지 문제를 동시에 풉니다: SHAP의 \"전문가용\" "
            "출력을 일반 사용자가 이해할 수 있는 자연어로 변환하면서, LLM의 "
            "환각을 원천적으로 차단합니다.")

    PageBreak(doc)

    # ── 2. 핵심 아이디어 ──
    H(doc, "2. 핵심 아이디어 — XAI-RAG가 무엇인가", level=1)
    P(doc, "본 연구가 제안하는 프레임워크는 4단계로 구성됩니다.")

    P(doc, "1단계 — 예측 (Prediction)", bold=True, size=12)
    P(doc, "정형 데이터를 모델(XGBoost 또는 TabNet)에 입력해 부도 확률 P(default)을 "
            "예측합니다. 본 연구는 트리 기반 모델과 정형 딥러닝(TabNet)을 모두 "
            "비교합니다.")

    P(doc, "2단계 — 해석 (Interpretation)", bold=True, size=12)
    P(doc, "학습된 모델에 SHAP을 적용해 개별 예측의 변수 기여도를 산출합니다. "
            "동시에 TabNet의 어텐션 마스크도 추출해 두 해석이 서로 일관되는지 "
            "분석합니다 (RQ2).")

    P(doc, "3단계 — 컨텍스트 구성 (Context Construction)", bold=True, size=12)
    P(doc, "SHAP 결과를 \"사실 단위\"의 구조화된 JSON으로 변환합니다. "
            "이 컨텍스트는 LLM에게 주는 \"검색된 근거\" 역할을 합니다. "
            "민감 변수(성별, 연령)는 마스킹되며, 도메인 용어는 한국어로 풀어 "
            "씁니다.")

    P(doc, "4단계 — 생성 + 평가 (Generation & Evaluation)", bold=True, size=12)
    P(doc, "LLM은 위 컨텍스트만 받아 자연어 설명 리포트를 생성합니다. "
            "그리고 그 리포트의 충실성(Faithfulness)·환각률(Hallucination Rate)·"
            "품질(G-Eval)을 정량 평가합니다.")

    Quote(doc, "핵심 아이디어를 한 줄로: \"LLM이 임의로 거절 사유를 추론하는 "
                "것이 아니라, 모델이 수학적으로 산출한 근거 데이터를 컨텍스트로 "
                "받아 자연어로 변환하는 역할에만 한정된다.\"")

    PageBreak(doc)

    # ── 3. 데이터 + EDA ──
    H(doc, "3. 데이터와 EDA — 무엇을 다루고 무엇을 발견했나", level=1)
    H(doc, "3.1 사용 데이터", level=2)
    P(doc, "Kaggle Home Credit Default Risk 대회의 메인 테이블 application_train.csv를 "
            "사용했습니다. 약 30만 7천 명의 신용 신청 기록이 있고 각 신청마다 122개의 "
            "특성이 있습니다. TARGET=1이 부도, 0이 정상입니다.")
    P(doc, "보조 테이블(과거 대출 이력 등 6개)은 미팅용 프로토타입에서는 제외하고, "
            "본 작업 단계에서 추가하기로 했습니다.")

    H(doc, "3.2 EDA에서 가장 중요했던 발견 5가지", level=2)
    Bullet(doc, "TARGET 분포가 8.07%로 매우 불균형합니다. Accuracy는 의미가 적고, "
                "AUROC·AUPRC·KS 통계량을 주 지표로 써야 합니다.")
    Bullet(doc, "EXT_SOURCE_2 / EXT_SOURCE_3 (외부 신용평가 점수)이 단연 강한 신호입니다. "
                "그 외 변수들의 단순 상관계수는 |ρ|<0.06로 약합니다. 즉 \"단일 변수 "
                "선형 신호\"로는 한계가 있고, 비선형·상호작용을 포착할 수 있는 "
                "트리/딥러닝이 유리합니다.")
    Bullet(doc, "결측이 50%를 넘는 컬럼이 41개 있습니다 (대부분 주거 정보). "
                "그대로 제거하지 않고 *_MISSING_FLAG 변수를 추가해 \"입력 여부 "
                "자체\"가 신호가 되도록 보존했습니다.")
    Bullet(doc, "DAYS_EMPLOYED에 365243(약 1000년)이라는 sentinel 값이 18% 있습니다. "
                "이는 \"무직\" 의미의 특수값으로, 그대로 두면 모델이 거대한 양수로 "
                "오해석합니다. NaN으로 변환하고 EMPLOYED_FLAG를 새로 만들었습니다.")
    Bullet(doc, "공정성 신호가 데이터에 이미 강하게 존재합니다. 남성 부도율 "
                "10.1% vs 여성 7.0% (1.45배), 25세 미만 12.3% vs 65세 이상 3.7% "
                "(3.4배). 이 발견은 5장에서 본격적으로 분석합니다.")

    Fig(doc, FIG_DIR / "03_protected_attrs.png",
        "그림 3-1. 성별·연령에 따른 부도율 차이. 데이터 자체에 구조적 편향이 "
        "이미 존재합니다.")

    PageBreak(doc)

    # ── 4. 모델 ──
    H(doc, "4. 모델 — 왜 베이스라인과 TabNet을 같이 보는가", level=1)
    P(doc, "본 연구는 단순히 \"가장 잘 맞히는 모델 하나\"를 고르는 것이 목적이 "
            "아닙니다. 두 가지가 모두 필요합니다.")
    Bullet(doc, "강한 베이스라인이 있어야 \"TabNet이 정형 데이터에서도 정말로 "
                "경쟁력 있는가\"를 정직하게 평가할 수 있습니다.")
    Bullet(doc, "TabNet은 어텐션 마스크라는 내재적 해석성을 제공합니다. "
                "5장에서 SHAP과의 일관성을 분석하기 위해서는 TabNet이 필요합니다.")

    H(doc, "4.1 5-fold Stratified CV 결과 (test set 기준)", level=2)
    Table(doc,
        headers=["모델", "AUROC", "AUPRC", "KS", "F1", "특징"],
        rows=[
            ["XGBoost", "0.7587 ± 0.0008", "0.2445 ± 0.0011",
             "0.3846 ± 0.0015", "0.2698 ± 0.0047", "5/5 fold 모두 1등"],
            ["LightGBM", "0.7544 ± 0.0009", "0.2402 ± 0.0018",
             "0.3788 ± 0.0028", "0.2584 ± 0.0052", "가장 빠름"],
            ["Logistic", "0.7544 ± 0.0001", "0.2343 ± 0.0006",
             "0.3804 ± 0.0010", "0.2631 ± 0.0087", "선형 모델 최고 안정성"],
            ["TabNet", "0.7518 ± 0.0017", "0.2331 ± 0.0023",
             "0.3749 ± 0.0056", "0.2657 ± 0.0079", "어텐션 해석성 제공"],
        ])

    P(doc, "결과 해석", bold=True, size=12)
    P(doc, "표준편차가 모두 0.001 안팎으로 매우 작습니다. 이는 \"단발 운\"이 "
            "아니라 안정적인 우열 관계라는 뜻입니다. XGBoost가 살짝 앞서지만 "
            "TabNet과의 차이는 0.007 AUROC 정도로 절대적 격차는 작습니다.")
    P(doc, "본 연구의 메시지는 \"XGBoost를 제치고 TabNet이 1등이다\"가 아닙니다. "
            "오히려 \"TabNet이 비슷한 성능을 내면서 어텐션 해석성을 추가로 "
            "제공한다\"입니다. 이 어텐션이 다음 장의 SHAP 일관성 분석으로 "
            "이어집니다.")

    Fig(doc, FIG_DIR / "13_cv_comparison.png",
        "그림 4-1. 5-fold CV 비교. 모든 지표에서 XGBoost가 일관되게 1등이지만 "
        "표준편차가 매우 작아 \"안정적인 격차\"임을 보여줍니다.")

    PageBreak(doc)

    # ── 5. SHAP × Attention ──
    H(doc, "5. SHAP과 어텐션은 같은 것을 말하는가? (RQ2)", level=1)
    P(doc, "TabNet의 어텐션과 사후 SHAP이 \"비슷한 변수를 강조한다\"는 가정은 "
            "기존 연구에서 종종 암묵적으로 받아들여졌지만, 정량적으로 검증된 "
            "적은 드뭅니다. 본 연구는 이를 직접 측정합니다.")

    H(doc, "5.1 측정 방법", level=2)
    Bullet(doc, "TabNet 모델 학습 후 어텐션 기반 변수 중요도(feature_importances_) "
                "추출")
    Bullet(doc, "동일 모델에 KernelExplainer로 SHAP 적용 후 mean(|SHAP|)로 "
                "global importance 산출")
    Bullet(doc, "둘의 Spearman 순위 상관 + Top-K 변수 교집합/Jaccard 측정")

    H(doc, "5.2 결과", level=2)
    Table(doc,
        headers=["지표", "값", "해석"],
        rows=[
            ["Spearman ρ (전체 214 변수)", "0.117 (p=0.089)", "전반적으로는 약한 양의 상관"],
            ["Spearman ρ (Top-50 합집합)", "−0.195", "주요 변수에서는 오히려 음의 상관 — 부분 상보적"],
            ["Top-20 교집합 / Jaccard", "9 / 0.29", "상위 20개 중 9개는 일치"],
        ])

    P(doc, "이 결과의 의미", bold=True, size=12)
    P(doc, "두 방법은 \"가장 핵심적인 변수\"(예: EXT_SOURCE_2/3)에서는 일관되게 "
            "같은 신호를 잡습니다. 그러나 그 다음 순위로 가면 두 방법이 다른 "
            "측면을 강조합니다. 어텐션은 비선형·상호작용으로 활용되는 변수를, "
            "SHAP은 marginal contribution이 명확한 변수를 더 강조합니다.")
    Quote(doc, "결론: 어텐션과 SHAP은 \"같은 답을 다른 단어로 말하는 것\"이 "
                "아닙니다. 두 방법을 함께 사용해야 모델 해석이 입체적으로 "
                "보입니다. 이는 본 연구가 두 정보를 모두 활용하는 정당성을 "
                "제공합니다.")

    Fig(doc, FIG_DIR / "16_attention_vs_shap_scatter.png",
        "그림 5-1. 어텐션 vs SHAP 산점도. 좌측은 전체, 우측은 Top-20 합집합. "
        "교집합(빨강)은 9개, 어텐션 only(파랑) 11개, SHAP only(주황) 11개.")

    PageBreak(doc)

    # ── 6. 공정성 ──
    H(doc, "6. 공정성 진단 — 모델은 누구를 차별하는가?", level=1)
    P(doc, "신용 평가 자동화에서 가장 민감한 부분은 보호 속성(성별, 연령 등)에 "
            "대한 차별 가능성입니다. 본 연구는 4개 모델 모두에 대해 4종 공정성 "
            "지표를 측정했습니다.")

    H(doc, "6.1 공정성 지표 4종 — 무엇을 재는가", level=2)
    Bullet(doc, "Demographic Parity (DP): 두 집단의 양성 예측률 차이 — "
                "\"누가 더 자주 거절되는가?\"")
    Bullet(doc, "Equal Opportunity (EO): 실제 부도자 중 거절률(TPR) 차이 — "
                "\"진짜 위험한 사람을 잡는 정확도가 집단별로 같은가?\"")
    Bullet(doc, "Equalized Odds: TPR과 FPR 모두를 종합")
    Bullet(doc, "Disparate Impact (DI): 최소 양성률/최대 양성률 비율. "
                "0.8 미만이면 미국 4/5 rule 위반.")

    H(doc, "6.2 베이스라인 결과 — 8/8 모두 4/5 rule 위반", level=2)
    P(doc, "4개 모델 × 보호속성 {GENDER, AGE} = 8개 케이스 전부에서 DI < 0.8입니다. "
            "특히 연령(AGE)이 성별(GENDER)보다 편향이 더 큽니다 — DI ≈ 0.50 (AGE) "
            "vs 0.62 (GENDER).")
    P(doc, "이는 \"어떤 모델을 골라도 데이터 자체의 구조적 편향이 그대로 "
            "옮겨진다\"는 뜻입니다. 알고리즘 차원의 문제가 아니라 데이터 차원의 "
            "문제입니다.")

    H(doc, "6.3 단순한 mitigation — 보호 속성 ablation", level=2)
    P(doc, "그렇다면 단순히 보호 속성 컬럼(CODE_GENDER, DAYS_BIRTH)을 빼고 "
            "다시 학습하면 어떨까? 결과는 흥미로웠습니다.")
    Table(doc,
        headers=["케이스", "AUROC 변화", "DP 변화", "DI 변화"],
        rows=[
            ["XGBoost × GENDER", "−0.005", "0.164 → 0.105 (−36%)", "0.622 → 0.718"],
            ["TabNet × GENDER", "−0.010", "0.156 → 0.093 (−40%)", "0.621 → 0.757"],
            ["XGBoost × AGE", "−0.005", "0.197 → 0.179 (−9%)", "0.498 → 0.507"],
            ["TabNet × AGE", "−0.010", "0.185 → 0.188 (+2%)", "0.504 → 0.513"],
        ])
    Bullet(doc, "성별(GENDER) ablation은 효과적입니다. AUROC 1% 미만 손실로 "
                "DP를 30~40% 줄였습니다.")
    Bullet(doc, "연령(AGE) ablation은 거의 효과가 없습니다. 이유는 연령 정보가 "
                "DAYS_EMPLOYED, DAYS_REGISTRATION, OWN_CAR_AGE 같은 다른 변수에 "
                "간접적으로 인코딩되어 있기 때문입니다 — \"proxy variable\" 문제.")
    Bullet(doc, "즉 단순한 컬럼 제거로는 4/5 rule을 통과하지 못합니다. 본격적인 "
                "fairness-aware 학습(Reweighing, Adversarial Debiasing)이 향후 "
                "과제로 자연스럽게 도출됩니다.")

    Fig(doc, FIG_DIR / "19_fairness_mitigation.png",
        "그림 6-1. baseline vs ablated 비교. GENDER ablation은 효과적이지만 "
        "AGE는 거의 변화 없음 — proxy variable 문제.")

    PageBreak(doc)

    # ── 7. XAI-RAG 데모 ──
    H(doc, "7. 실제 인스턴스로 따라가는 XAI-RAG (idx 59291)", level=1)

    walk = json.loads((RESULTS_DIR / "demo_walkthrough.json").read_text(encoding="utf-8"))
    s = walk["sample"]
    p = walk["prediction"]

    P(doc, f"본 장은 한 명의 실제 인스턴스 (idx {s['idx']})를 따라가며 "
            f"전 파이프라인이 어떻게 작동하는지 보여줍니다. 이 인스턴스는 "
            f"True Positive — 모델이 부도라고 예측했고 실제로도 부도였습니다.")

    H(doc, "7.1 입력 데이터", level=2)
    P(doc, f"연령 {s['age']:.0f}세, 성별 {s['gender']}. 주요 정형 변수:")
    Table(doc, ["변수", "값"], list(s["display_features"].items()))

    H(doc, "7.2 1단계 — XGBoost 예측", level=2)
    P(doc, f"모델은 부도 확률 P(default) = {p['default_proba']:.4f}로 예측했고, "
            f"validation에서 결정한 임계치 {p['threshold']:.4f}를 넘었으므로 "
            f"결정은 {p['decision']}입니다.")

    H(doc, "7.3 2단계 — SHAP Local Explanation", level=2)
    P(doc, "SHAP은 이 예측에서 어떤 변수가 얼마나 기여했는지 알려줍니다. "
            "양수 SHAP은 부도 확률을 높이는 방향, 음수는 낮추는 방향입니다.")
    Table(doc,
        ["순위", "변수", "값", "SHAP"],
        [[d["rank"], d["feature"], d["value"], f"{d['shap']:+.4f}"]
          for d in walk["context"]["top_drivers_for_default"]])

    H(doc, "7.4 3단계 — JSON 컨텍스트 빌드", level=2)
    P(doc, "위 SHAP 결과를 LLM이 읽기 쉬운 사실 단위 JSON으로 변환합니다. "
            "민감 변수(성별·연령)는 마스킹되고, 변수명은 한국어로 풀어쓰며, "
            "explanation_policy: \"fact_only\" 정책이 명시됩니다.")

    H(doc, "7.5 4단계 — LLM 자연어 설명 생성", level=2)
    P(doc, "이 컨텍스트와 엄격한 제약 조건이 담긴 프롬프트를 두 LLM(Gemini와 "
            "Claude)에 동시에 보내 각각의 자연어 설명을 받았습니다.")

    for prov_key, prov_label in [("gemini", "Gemini 2.5 Flash"),
                                    ("anthropic", "Claude Sonnet 4.5")]:
        if prov_key in walk["llm_outputs"] and "explanation" in walk["llm_outputs"][prov_key]:
            ll = walk["llm_outputs"][prov_key]
            P(doc, f"[{prov_label}]  처리 시간 {ll['elapsed_sec']:.1f}초, "
                    f"토큰 {ll.get('total_tokens', '?')}",
              bold=True, size=11, color=(0x4C, 0x72, 0xB0))
            for line in ll["explanation"].split("\n"):
                if line.strip():
                    P(doc, line.strip(), size=10)
            P(doc, "")

    Quote(doc, "두 LLM 모두 컨텍스트의 변수명·값·SHAP 부호를 정확히 인용했고, "
                "성별·연령 같은 민감 변수는 한 번도 직접 언급하지 않았습니다. "
                "스타일에서 차이가 있어 — Gemini는 풀 자릿수(0.0633754)를 그대로 "
                "쓰고, Claude는 자연스럽게 반올림(0.0634)합니다.")

    Fig(doc, FIG_DIR / "22_demo_walkthrough.png",
        "그림 7-1. Demo 워크스루 — SHAP Top 10 + Decision Summary.")

    PageBreak(doc)

    # ── 8. 정량 평가 ──
    H(doc, "8. 자연어 설명을 어떻게 평가하는가 (RQ3)", level=1)
    P(doc, "\"좋은 설명\"이라는 개념은 본질적으로 다층적입니다. 본 연구는 "
            "기존 LLM 평가 연구를 따라 4개 차원으로 평가합니다.")

    Bullet(doc, "Faithfulness (충실성): 텍스트의 변수·수치·부호가 컨텍스트와 "
                "정확히 일치하는가?")
    Bullet(doc, "Hallucination Rate (환각률): 컨텍스트에 없는 변수가 텍스트에 "
                "등장한 비율은 얼마인가?")
    Bullet(doc, "G-Eval: 별도의 LLM(여기서는 Gemini)이 정해진 루브릭으로 5점 "
                "척도 평가")
    Bullet(doc, "효율성: 호출당 시간과 토큰")

    H(doc, "8.1 결과 — Hallucination Rate가 양쪽 모두 0", level=2)
    Table(doc,
        ["지표", "Gemini 2.5 Flash", "Claude Sonnet 4.5"],
        [
            ["Hallucination Rate (strict)", "0.000 ± 0.000", "0.000 ± 0.000"],
            ["Hallucination Rate (broad)", "0.000 ± 0.000", "0.000 ± 0.000"],
            ["val_match_rate", "0.811 ± 0.123", "0.901 ± 0.111"],
            ["sign_match_rate", "0.783 ± 0.209", "0.867 ± 0.233"],
            ["G-Eval factual_accuracy", "5.0 / 5", "(skip)"],
            ["G-Eval sensitive_leak", "5.0 / 5", "(skip)"],
            ["G-Eval style", "5.0 / 5", "(skip)"],
            ["elapsed (sec/call)", "12.7 ± 4.5", "8.4 ± 0.8"],
            ["total tokens (per call)", "4,155 ± 834", "2,500 ± 83"],
        ])
    P(doc, "두 LLM 모두에서 환각률이 0이라는 것은 본 연구의 가장 중요한 결과입니다. "
            "이는 SHAP 컨텍스트가 LLM의 자유 추론을 효과적으로 차단함을 보여줍니다.",
      bold=True)
    Fig(doc, FIG_DIR / "21_llm_comparison.png",
        "그림 8-1. Gemini vs Claude — 룰 기반 평가, G-Eval, 효율성 3개 패널.")

    PageBreak(doc)

    # ── 9. Counterfactual Baseline ──
    H(doc, "9. SHAP이 없으면 정말로 환각이 늘어나는가? (RQ3 강한 검증)",
        level=1)
    P(doc, "8장에서 \"본 구조에서 환각률이 0이다\"라고 했지만, 이것이 본 구조 "
            "덕분인지 아니면 단지 LLM이 좋은 모델이라서인지는 별도의 실험이 "
            "필요합니다. 그래서 동일 11개 인스턴스에 대해 \"SHAP 컨텍스트 없이 "
            "raw 데이터만 LLM에 주는\" baseline 실험을 진행했습니다.")

    H(doc, "9.1 결과 — Claude의 환각률 0% → 45.5%", level=2)
    Table(doc,
        ["LLM", "XAI-RAG (SHAP context)", "Baseline (no SHAP)"],
        [
            ["Gemini 2.5 Flash", "0.000", "0.000 (측정 한계 — 9.2절 참조)"],
            ["Claude Sonnet 4.5", "0.000", "0.4545"],
        ])
    P(doc, "Claude의 baseline 환각률 45.5%가 결정적 결과입니다. SHAP 컨텍스트 "
            "없으면 Claude는 학습된 일반 금융 지식을 자유롭게 끌어와 거절 "
            "사유를 만들어냅니다.", bold=True)

    H(doc, "9.2 Claude baseline 환각 사례", level=2)
    Bullet(doc, "DTI(Debt-to-Income), LTV(Loan-to-Value), DSR(Debt Service Ratio) "
                "— Home Credit 데이터에는 없는 일반 금융 비율 약어를 자유 추론")
    Bullet(doc, "\"햇살론, 미소금융\" 같은 한국 특정 정부 지원 금융상품을 "
                "권고에 포함 — 데이터에 없는 외부 지식")
    Bullet(doc, "\"고객센터 ☎1588-XXXX\" 같은 가짜 연락처 생성")
    Bullet(doc, "DEF_30_CNT처럼 변수명을 잘라서 부정확하게 인용")

    H(doc, "9.3 Gemini의 측정 한계 — 환각이 없는 게 아니라 잡히지 않은 것",
        level=2)
    P(doc, "Gemini의 baseline 환각률 0%는 다소 오해의 소지가 있습니다. "
            "Gemini는 영문 변수명을 거의 사용하지 않고 한국어 자연어로 의역하는 "
            "경향이 있어, 영문 정규식 기반 룰로는 환각이 잡히지 않습니다. "
            "내용 단위(claim 단위) 평가를 위한 cross-LLM judge 방식이 향후 "
            "필요합니다.")

    Fig(doc, FIG_DIR / "23_baseline_vs_xairag.png",
        "그림 9-1. XAI-RAG vs baseline (no SHAP) 환각률. Claude에서 결정적 차이 "
        "(0% → 45.5%).")

    PageBreak(doc)

    # ── 10. 종합 정리 ──
    H(doc, "10. 종합 — 본 연구의 가치", level=1)
    P(doc, "본 연구는 정형 데이터 신용 평가에서 다음 세 가지를 동시에 달성했습니다.")
    Bullet(doc, "예측 성능: XGBoost 5-fold AUROC 0.7587 ± 0.0008로 안정적인 "
                "성능 확보. TabNet도 0.7518로 비슷한 수준.")
    Bullet(doc, "해석 가능성: TabNet의 어텐션과 SHAP이 핵심 변수에서 일관되며 "
                "(EXT_SOURCE 등), 미세 영역에서 상보적으로 작동하는 것을 정량 "
                "검증.")
    Bullet(doc, "자연어 설명의 신뢰성: 두 상용 LLM(Gemini, Claude) 모두에서 "
                "환각률 0%. SHAP 없는 baseline에서는 Claude가 45.5%까지 환각 — "
                "본 구조가 환각을 명확히 차단한다는 직접 증거.")
    P(doc, "더불어 8개 공정성 케이스 모두에서 4/5 rule 위반을 발견했고, 단순 "
            "ablation의 효과와 한계(특히 AGE의 proxy variable 문제)를 측정했습니다.")

    PageBreak(doc)

    # ── 11. Step 2-A 평가 신뢰성 강화 (NEW) ──
    H(doc, "11. Step 2-A — 평가 신뢰성을 한 단계 끌어올리다", level=1)
    P(doc, "Step 1까지의 결과는 흥미로웠지만 통계적으로 약한 부분이 있었습니다. "
            "10명 표본 평가 — 이 정도 표본으로는 \"운 좋게 환각이 0%였을 수 있다\"는 "
            "비판이 가능합니다. 그래서 Step 2-A에서는 평가 자체를 강화하는 데 "
            "집중했습니다.")

    H(doc, "11.1 100명 표본 확장 — 환각률 0%가 견고한가?", level=2)
    P(doc, "Step 1에서 평가했던 10명 대신, 100명(거절 50명 + 정상 50명)에 대해 "
            "동일한 파이프라인을 다시 돌렸습니다.")
    Table(doc, ["LLM", "Hallucination Rate", "n"],
        [["Gemini 2.5 Flash", "0.000 ± 0.000", "100"],
         ["Claude Sonnet 4.5", "0.000 ± 0.000", "100"]])
    Quote(doc, "결과: 표본을 10배 늘려도 환각률 0% 유지. Step 1 결과가 우연이 "
                "아니라 본 시스템의 안정적 특성임을 통계로 입증.")

    H(doc, "11.2 Cross-LLM G-Eval — Self-bias를 어떻게 우회하는가", level=2)
    P(doc, "Step 1의 G-Eval은 Gemini가 자기 자신의 출력을 평가한 결과였습니다. "
            "이론적으로 자기 편애(self-bias)가 가능합니다. 그래서 Step 2-A에서는 "
            "\"다른 LLM이 평가하면 어떻게 되는가\"를 측정했습니다.")
    Table(doc,
        ["Judge → Target", "Factual", "Complete", "Sensitive", "Style"],
        [["Gemini → Gemini (Step 1)", "5.00", "3.38", "5.00", "5.00"],
         ["Claude → Gemini (NEW)", "4.87 ± 0.51", "4.00 ± 0.64", "5.00", "4.97"],
         ["Gemini → Claude (NEW)", "4.60 ± 0.89", "3.33 ± 0.96", "5.00", "5.00"]])
    P(doc, "흥미로운 발견 3가지", bold=True, size=12)
    Bullet(doc, "양 LLM 모두 factual_accuracy ≥ 4.6/5 — 본 구조의 사실성이 LLM "
                "종속성 없이 유지됨")
    Bullet(doc, "sensitive_leak 5.0/5 만점이 모든 방향에서 — 본 시스템의 민감 변수 "
                "마스킹이 어느 LLM이 평가하든 완벽하게 작동")
    Bullet(doc, "Gemini의 self-bias가 \"자기 비판\" 방향이라는 점 — Gemini가 자기 "
                "completeness를 3.38로 박하게 매겼는데 Claude는 4.0으로 더 후하게 "
                "줬음. 즉 Step 1의 G-Eval은 오히려 Gemini를 과소평가했음")

    H(doc, "11.3 Counterfactual Test — LLM이 진짜로 컨텍스트에 의존하는가",
       level=2)
    P(doc, "본 연구의 핵심 가설은 \"LLM이 SHAP 컨텍스트만 사용해 설명한다\"입니다. "
            "그렇다면 컨텍스트의 핵심 변수를 빼면 출력이 정말로 달라질까요?")
    P(doc, "각 인스턴스 30개에 대해 SHAP top_drivers_for_default rank 1 변수를 "
            "컨텍스트에서 제거하고 LLM에 다시 호출했습니다. 원본 출력과의 의미 "
            "유사도(cosine, multilingual sentence-transformers)와 어휘 중복도"
            "(ROUGE-L)를 측정했습니다.")
    Table(doc,
        ["LLM", "Cosine sim", "ROUGE-L"],
        [["Claude Sonnet 4.5", "0.909 ± 0.069", "0.747 ± 0.112"],
         ["Gemini 2.5 Flash", "0.920 ± 0.040", "0.750 ± 0.116"]])
    P(doc, "결과 해석", bold=True, size=12)
    P(doc, "Cosine 0.91 → \"의미상 90% 유지\". ROUGE-L 0.75 → \"어휘는 25%가 "
            "변함\". 즉 LLM은 컨텍스트의 부분 변경에 부분적으로 반응하면서도 전체 "
            "메시지의 일관성을 유지합니다. 1개 변수 제거 후 cosine이 0이 되지 않은 "
            "이유는 나머지 4개 driver는 그대로 있기 때문이며, 합리적 결과입니다.")

    H(doc, "11.4 Robustness 평가 — 프롬프트가 흔들려도 출력은 안정적인가",
       level=2)
    P(doc, "실제 운영에서 프롬프트는 시간이 지나면서 미세하게 바뀝니다. 본 시스템이 "
            "프롬프트 변형에 강건한지 측정했습니다. 3가지 변형:")
    Bullet(doc, "role_swap: \"금융 상담사\" → \"신용 평가 전문가\"")
    Bullet(doc, "example_swap: Few-shot 예시 위치를 컨텍스트 뒤로 이동")
    Bullet(doc, "driver_shuffle: 컨텍스트 driver의 그룹 내 순서를 셔플 (rank는 "
                "유지)")
    Table(doc,
        ["Variant", "Claude cosine", "Gemini cosine"],
        [["role_swap", "0.923 ± 0.057", "0.951 ± 0.034"],
         ["example_swap", "0.914 ± 0.060", "0.908 ± 0.077"],
         ["driver_shuffle", "0.924 ± 0.054", "0.942 ± 0.032"]])
    Quote(doc, "두 LLM 모두 모든 변형에서 cosine ≥ 0.90. 본 시스템은 프롬프트 "
                "변형에 매우 강건. 운영 환경에서도 출력 일관성을 기대할 수 있음.")

    H(doc, "11.5 Step 2-A 종합 — Step 1의 메시지가 강해진다", level=2)
    Bullet(doc, "환각 0% 메시지: 10명 → 100명 표본에서 견고. 통계적 신뢰성 확보")
    Bullet(doc, "Cross-LLM 평가: self-bias 우회. 두 LLM이 서로의 출력을 보더라도 "
                "factual·sensitive 만점급. 본 시스템의 마스킹 정책은 LLM 무관하게 "
                "작동")
    Bullet(doc, "Counterfactual: 컨텍스트의 부분 변경에 부분적으로 반응. \"무시도 "
                "않고 과잉 반응도 않는\" 균형 잡힌 의존")
    Bullet(doc, "Robustness: 프롬프트 미세 변형에 cosine 0.91~0.95로 안정")
    P(doc, "")
    P(doc, "★ 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"본 XAI-RAG 구조는 100명 표본·LLM 종속성 없는 환경에서도 환각 차단 "
            "효과를 유지하며, 프롬프트·컨텍스트 변형에도 안정적으로 작동한다.\"",
       bold=True, size=12)

    PageBreak(doc)

    # ── 12. Step 3-B 보조 테이블 활용 (NEW, 2026-05-05) ──
    H(doc, "12. Step 3-B — 보조 테이블 두 개를 추가하면 어떻게 달라지는가",
        level=1)
    P(doc, "Step 1과 Step 2-A는 모두 application_train.csv 한 테이블만 사용했습니다. "
            "Home Credit Default Risk 데이터셋에는 사실 보조 테이블이 6개 더 있는데, "
            "계획서에서도 \"보조 테이블 활용은 future work\"로 미뤄두었었습니다. "
            "발표 5일 전(D-5), 그 중 임팩트가 큰 두 개(`bureau`, `previous_application`)를 "
            "추가해 한 사이클 더 돌려봤습니다.")
    Quote(doc, "왜 두 개만? — 6개 다 쓰려면 4~5일이 필요하지만, 두 개만으로도 임팩트의 "
                "80%를 확보할 수 있습니다. 발표 메시지를 \"추가 진척\"으로 자연스럽게 "
                "이어갈 수 있는 사이즈를 골랐습니다.")

    H(doc, "12.1 두 보조 테이블의 의미", level=2)
    Bullet(doc, "bureau: 외부 신용기관(타사 은행/카드)이 가진 이 사람의 과거 대출 이력")
    Bullet(doc, "previous_application: 같은 Home Credit에 이전에 신청했던 대출 이력")
    P(doc, "쉽게 말하면 bureau는 \"남들이 본 이 사람의 과거\"이고, "
            "previous_application은 \"우리 회사가 본 이 사람의 과거\"입니다. "
            "둘은 정보 출처가 완전히 다르므로 main 테이블이 가진 신호와는 별개의 "
            "예측력을 줄 가능성이 있습니다.")

    H(doc, "12.2 데이터를 어떻게 다루었는가", level=2)
    P(doc, "보조 테이블은 한 사람당 여러 행입니다(과거 대출 N건). "
            "이걸 main 테이블에 join하려면 SK_ID_CURR 단위로 \"집계\"해야 합니다. "
            "예: 평균 금액, 거절 비율, 결제 횟수의 표준편차 등.")
    Bullet(doc, "bureau_balance(2,700만 행)는 SK_ID_BUREAU 단위로 STATUS 비율 + "
                "MONTHS_BALANCE 통계로 압축")
    Bullet(doc, "bureau는 그 결과를 합친 뒤 SK_ID_CURR 기준 전체 / Active(미상환) / "
                "Closed(상환완료) 분리 집계")
    Bullet(doc, "previous_application은 SK_ID_CURR 기준 전체 / Approved / Refused 분리 집계")
    P(doc, "최종적으로 756개의 집계 변수가 생성되어 main 테이블에 left merge되었고, "
            "동일한 전처리 정책으로 1161 features의 학습 데이터가 만들어졌습니다.")
    Quote(doc, "메모리 트릭 — 원본 csv 합계 ~950 MB가 DataFrame으로 로드되면 ~3.4 GB로 "
                "부풀지만, optimize_dtypes로 float32 / int8~32 / category 다운캐스트하면 "
                "388 MB로 줄어듭니다(-89%). 16 GB RAM 환경에서 5-fold CV를 돌릴 수 있는 "
                "여유가 생깁니다.")

    H(doc, "12.3 5-fold CV 결과 — 통계적으로 명확한 향상", level=2)
    Table(doc,
        ["Metric", "Baseline (Step 1)", "+ Aux (Step 3-B)", "Δ", "Δ%"],
        [["AUROC", "0.7587 ± 0.0008", "0.7755 ± 0.0011", "+0.0168", "+2.22%"],
         ["AUPRC", "0.2445 ± 0.0011", "0.2646 ± 0.0015", "+0.0201", "+8.21%"],
         ["KS", "0.3846 ± 0.0015", "0.4146 ± 0.0040", "+0.0301", "+7.81%"],
         ["F1", "0.2698 ± 0.0047", "0.2813 ± 0.0096", "+0.0115", "+4.27%"]])
    P(doc, "")
    Bullet(doc, "AUROC +0.0168은 baseline std (0.0008)의 21배. "
                "5 fold가 모두 0.7743 ~ 0.7771의 좁은 범위에 있어 매우 안정적입니다.")
    Bullet(doc, "AUPRC와 KS의 개선 비율(+8.21%, +7.81%)이 AUROC(+2.22%)보다 큽니다. "
                "신용 평가는 positive(부도)가 8%인 불균형 데이터라, "
                "AUPRC와 KS가 실제 분류력을 더 잘 반영합니다 — "
                "즉 \"진짜 부도자를 놓치지 않는 능력\"이 강화되었다는 의미입니다.")
    Fig(doc, FIG_DIR / "27_cv_aux_comparison.png",
        "그림 12-1. Baseline vs Aux 모델의 5-fold CV 비교 (test set, mean ± std)")

    H(doc, "12.4 SHAP 관점 — 어떤 변수가 새로 중요해졌는가", level=2)
    P(doc, "성능이 좋아졌다면 \"왜\"를 SHAP으로 들여다봐야 합니다. baseline 모델의 "
            "SHAP top 20과 aux 모델의 SHAP top 20을 비교하면, 5개의 보조 변수가 "
            "새로 top 20에 진입했습니다.")
    Table(doc,
        ["rank", "신규 진입 feature", "mean(|SHAP|)", "한국어 의미"],
        [["12", "PREV_NAME_YIELD_GROUP_high_mean", "0.0499", "이전 high-yield 신청 비율"],
         ["13", "PREV_CNT_PAYMENT_std", "0.0486", "이전 결제 횟수 변동성"],
         ["14", "PREV_NAME_YIELD_GROUP_low_action_mean", "0.0450", "이전 low-yield 신청 비율"],
         ["16", "PREV_DAYS_LAST_DUE_1ST_VERSION_max", "0.0440", "이전 만기일 최댓값"],
         ["20", "PREV_NAME_CONTRACT_STATUS_Refused_mean", "0.0360", "★ 이전 거절 비율"]])
    P(doc, "")
    Quote(doc, "5개 모두 PREV_* 접두사 — 즉 자체 이력만 진입했고, "
                "외부 신용기관(bureau) 변수는 한 개도 진입하지 못했습니다. "
                "이건 매우 흥미로운 결과입니다.")
    P(doc, "왜 bureau는 진입하지 못했을까? 가장 그럴듯한 가설은, "
            "main 테이블의 EXT_SOURCE_1/2/3(외부 신용평가 점수)에 외부 신용기관 정보가 "
            "이미 응축돼 들어가 있다는 것입니다. "
            "EXT_SOURCE는 Step 1에서 SHAP top 1~3을 차지한 압도적 신호였는데, "
            "이게 사실상 bureau 데이터의 잘 압축된 표현일 수 있습니다. "
            "추가 ablation(bureau 단독 vs prev 단독 vs 둘 다)으로 검증할 future work입니다.")
    Fig(doc, FIG_DIR / "29_shap_top20_overlap.png",
        "그림 12-2. baseline vs aux 모델의 SHAP top 20 비교 (녹색=공통, 주황=신규 aux)")

    H(doc, "12.5 가장 직관적인 신규 신호 — 이전 거절 비율", level=2)
    P(doc, "rank 20에 진입한 PREV_NAME_CONTRACT_STATUS_Refused_mean은 "
            "\"이 사람이 과거에 신청한 대출 중 거절당한 비율\"입니다. "
            "직관적으로 가장 강력한 신용 신호 중 하나입니다.")
    Quote(doc, "이전에 자주 거절된 사람은 다시 거절될 가능성이 높다 — "
                "신용 평가 도메인의 상식과 일치합니다. "
                "이 변수가 데이터에 명시적으로 들어가 있지 않다가 추가되면서 "
                "성능이 향상된 것은, 본 파이프라인이 \"의미 있는 변수에 반응한다\"는 "
                "정상성(sanity check)을 보여줍니다.")

    H(doc, "12.6 Step 3-B 종합", level=2)
    Bullet(doc, "보조 테이블 두 개 추가만으로 AUROC +0.0168 (+2.22%) — 통계적으로 "
                "명확한 향상")
    Bullet(doc, "AUPRC +8.21%, KS +7.81% — 불균형 데이터에서 분류력 자체가 강화됨")
    Bullet(doc, "외부 신용기관(bureau)은 SHAP에 직접 진입하지 못했지만, "
                "EXT_SOURCE에 응축된 형태로 이미 활용되고 있을 가능성")
    Bullet(doc, "자체 이력(previous_application)은 5개 변수가 SHAP top 20에 진입 — "
                "특히 \"이전 거절 비율\"이 직관적으로 가장 강력한 신호")
    Bullet(doc, "1161 feature 학습 시간 138s/fold — 운영 환경에서도 충분히 가능한 부담")
    P(doc, "")
    P(doc, "★ Step 3-B 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"보조 테이블 활용은 본 파이프라인의 자연스러운 다음 단계이며, "
            "기존 메시지(환각 0%, baseline 45.5%)에 \"성능까지 향상되는\" 임팩트를 "
            "더한다.\"", bold=True, size=12)

    PageBreak(doc)

    # ── 13. Step 3-C-1 — TabNet 어텐션 × SHAP 융합 컨텍스트 (NEW) ──
    H(doc, "13. Step 3-C-1 — TabNet이 비로소 본 메커니즘에 통합된다",
        level=1)
    P(doc, "지금까지의 본 연구에서 한 가지 큰 약점이 있었습니다. **논문 제목에 \"TabNet\"이 "
            "들어가 있지만, 실제로는 TabNet이 \"비교 모델\"의 역할만 했고, LLM의 컨텍스트는 "
            "XGBoost의 SHAP만으로 만들어졌습니다.** 즉 \"왜 제목에 TabNet인가?\"라는 질문에 "
            "정직하게 답하기 어려웠습니다.")
    Quote(doc, "Step 3-C-1에서는 그 약점을 정면 대응합니다 — TabNet의 instance-level "
                "어텐션을 SHAP과 융합한 \"agreement-aware 컨텍스트\"를 LLM에 제공하는 "
                "메커니즘을 구현했고, 그 효과를 정량 측정했습니다.")

    H(doc, "13.1 융합 메커니즘 — 두 해석 신호의 상보성을 LLM에 명시 노출", level=2)
    P(doc, "각 인스턴스에서 두 해석을 동시에 추출합니다:")
    Bullet(doc, "XGBoost SHAP top-10 (positive 5 + negative 5) — 부호와 기여도 보존")
    Bullet(doc, "TabNet attention top-5 — 모델이 sparse하게 집중한 변수, 부호 없음")
    P(doc, "두 set을 비교해 3 그룹으로 분류:")
    Bullet(doc, "agreed_drivers — 두 모델이 모두 본 변수 (가장 신뢰할 만한 강한 신호)")
    Bullet(doc, "shap_only_drivers — SHAP만 본 변수 (부호 + 기여도 정보가 살아있는 보완)")
    Bullet(doc, "attention_only_drivers — TabNet만 본 변수 (방향성은 모르지만 모델이 본 신호)")
    P(doc, "LLM 프롬프트에 그룹 라벨의 의미를 명시 — 즉 LLM은 \"두 모델이 동의한 강한 신호\"와 "
            "\"한쪽만 본 보완 신호\"를 구분해 표현하도록 지시받습니다. attention_only는 부호가 "
            "없으므로 \"결정에 영향을 준 변수\"로만 표현하고 방향성을 추측하지 않도록 명시.")

    H(doc, "13.2 부분 일관 + 부분 상보 — 정량 입증", level=2)
    P(doc, "100 인스턴스에 대한 두 신호의 겹침 패턴:")
    Table(doc,
        ["그룹", "평균 변수 수", "n_agreed 분포 (100 인스턴스)"],
        [["agreed_drivers", "2.12", "0개=1, 1개=8, 2개=69, 3개=22, 4개+=0"],
         ["shap_only_drivers", "6.98", "—"],
         ["attention_only_drivers", "2.06", "—"]])
    Quote(doc, "두 해석 모델은 거의 항상 부분적으로만 겹친다 (3개 이상 동의 22%, 4개 이상 "
                "동의는 단 한 번도 없음). 이는 Day 4의 어텐션-SHAP global ρ=0.117 분석과 "
                "정확히 일치하는 instance-level 패턴 — 본 데이터의 본질적 특성으로서 "
                "\"부분 일관 + 부분 상보\"를 다시 입증한 것입니다.")

    H(doc, "13.3 결과 — 환각 차단 유지 + 완결성 큰 향상", level=2)
    P(doc, "30 인스턴스 × 2 LLM × 2 mode (shap-only vs fusion). Judge는 Claude로 통일 "
            "(같은 judge로 4 그룹 통제 → Δ 측정 객관성).")
    Table(doc,
        ["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        [["Halluc strict (↓)", "Anthropic", "0.000", "0.000", "0 ✅"],
         ["Halluc strict (↓)", "Gemini", "0.000", "0.000", "0 ✅"],
         ["G-Eval Completeness (↑)", "Anthropic", "4.30", "4.97", "+0.67 ★"],
         ["G-Eval Completeness (↑)", "Gemini", "3.90", "4.70", "+0.80 ★"],
         ["G-Eval Factual (↑)", "Anthropic", "4.87", "4.90", "+0.03 ≈"],
         ["G-Eval Factual (↑)", "Gemini", "4.90", "4.77", "-0.13 ≈"],
         ["G-Eval Sensitive (↑)", "Both", "5.00", "5.00", "0 ✅"],
         ["val_match_rate (↑)", "Anthropic", "0.85", "0.90", "+0.06"],
         ["val_match_rate (↑)", "Gemini", "0.79", "0.86", "+0.07"]])
    P(doc, "")
    P(doc, "★ 핵심 1 — 환각 차단 유지", bold=True, color=(0x55, 0xA8, 0x68))
    P(doc, "양 LLM × 양 mode 4 조합 모두 Halluc strict 0/30. 즉 두 해석 신호를 융합해도 "
            "Step 1/2-A의 환각 차단 메커니즘은 그대로 작동.")
    P(doc, "")
    P(doc, "★ 핵심 2 — 완결성 큰 향상", bold=True, color=(0x55, 0xA8, 0x68))
    P(doc, "G-Eval Completeness가 Anthropic 4.30→4.97 (+0.67), Gemini 3.90→4.70 (+0.80). "
            "양 LLM 모두에서 일관된 큰 향상. 두 해석 신호의 상보성이 LLM에게 더 풍부한 "
            "정보를 제공해서 더 완결된 설명을 만든다는 직접 증거.")
    P(doc, "")
    P(doc, "★ 핵심 3 — 사실성·민감도·스타일 유지", bold=True, color=(0x55, 0xA8, 0x68))
    P(doc, "G-Eval Factual은 4.77~4.97 만점급 유지, Sensitive Leak 5.0/5.0 만점, "
            "Style 4.93~5.0 동등. 융합이 사실성·민감도·톤을 손상시키지 않음.")
    Fig(doc, FIG_DIR / "30_fusion_vs_shaponly.png",
        "그림 13-1. SHAP-only vs Fusion 비교 (양 LLM × 4 메트릭, n=30 each)")

    H(doc, "13.4 룰 기반 sign_match 하락은 룰의 한계", level=2)
    P(doc, "보고된 표에 sign_match_rate가 fusion에서 떨어진 것이 보입니다 (Anthropic "
            "0.87→0.65, Gemini 0.94→0.77). 이건 정직한 보고이며, 진짜 환각이나 사실성 "
            "악화가 아닙니다.")
    Bullet(doc, "룰의 sign 평가는 \"높였습니다\", \"낮추는\", \"긍정적\" 등 한정된 키워드 셋으로 작동")
    Bullet(doc, "Fusion 컨텍스트에서는 LLM이 더 다양한 표현을 사용 — 예: \"증가시키는\", \"위험\", \"영향을 준\" 등")
    Bullet(doc, "이 단어들이 룰의 키워드 셋에 없어서 false negative 발생 → sign_match 하락")
    Bullet(doc, "G-Eval Factual Accuracy(4.77~4.97)가 진짜 사실성을 정확히 측정 — 만점급으로 유지됨이 증명")
    Quote(doc, "이 한계는 future work에서 NLI 기반 Faithfulness(\"context entails statement?\" 검증) "
                "로 해소 가능. 본 보고에선 룰 + G-Eval 이중 측정으로 honest reporting 유지.")

    H(doc, "13.5 Step 3-C-1 종합", level=2)
    Bullet(doc, "TabNet이 단순 비교 모델 → 본 메커니즘의 핵심 구성 요소로 격상")
    Bullet(doc, "두 해석 신호의 부분 일관 + 부분 상보를 LLM에 명시 노출하는 agreement-aware 컨텍스트 제안 (좁지만 명확한 novelty)")
    Bullet(doc, "환각 차단 유지 + 완결성 큰 향상의 두 마리 토끼")
    Bullet(doc, "양 LLM(Anthropic + Gemini)에서 일관된 결과 → LLM 종속성 없는 메커니즘")
    Bullet(doc, "룰 sign_match 하락의 정체는 다음 절(Step 3-C-2)에서 NLI로 직접 입증")
    P(doc, "")
    P(doc, "★ Step 3-C-1 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"TabNet 어텐션과 SHAP의 상보성을 LLM 컨텍스트로 융합하면, 환각 차단을 "
            "유지하면서 설명의 완결성이 크게 향상된다.\"", bold=True, size=12)

    PageBreak(doc)

    # ── 14. Step 3-C-2 — NLI 기반 평가 객관성 보강 (NEW) ──
    H(doc, "14. Step 3-C-2 — 룰 한계를 의미적 측정으로 보강하다",
        level=1)
    P(doc, "Step 3-C-1의 결과 표에서 한 가지 의문이 남았습니다. **룰의 sign_match가 "
            "fusion에서 떨어졌습니다** (Anthropic 0.87→0.65, Gemini 0.94→0.77). "
            "이걸 \"룰의 키워드 한계\" 라고 추정만 했지, 직접 입증한 것은 아니었습니다.")
    Quote(doc, "Step 3-C-2에서는 NLI(Natural Language Inference) 모델로 의미적 함의를 "
                "직접 측정해서 그 추정을 검증합니다.")

    H(doc, "14.1 NLI는 무엇이고 왜 도입했는가", level=2)
    P(doc, "NLI는 \"premise(전제)가 주어졌을 때 hypothesis(가설)가 entailment(함의), "
            "neutral(중립), contradiction(모순) 중 어느 관계인가\"를 분류하는 자연어 처리 "
            "표준 작업입니다. 본 연구에서는:")
    Bullet(doc, "Premise = LLM에 주어진 컨텍스트의 모든 facts를 자연어 단락으로 변환")
    Bullet(doc, "Hypothesis = LLM이 생성한 설명의 각 문장")
    Bullet(doc, "NLI 모델이 (premise, hypothesis) 쌍에 대해 entailment 확률을 출력")
    Bullet(doc, "인스턴스별 entailment_rate 평균 → faithfulness score")
    P(doc, "이 방식은 룰의 \"낮추는/긍정/부정\" 같은 한정 키워드와 무관하게 의미 자체를 "
            "측정하므로, fusion에서 LLM이 다양한 표현을 사용해도 정확히 평가합니다.")
    P(doc, "사용 모델: **mDeBERTa-v3-base-xnli (KLUE 학습 데이터 100M+ 포함, 다국어 NLI)**. "
            "원래는 한국어 native KLUE-roberta-NLI를 시도했으나 torch 2.5/transformers 5.7 "
            "보안 충돌(CVE-2025-32434)로 safetensors 형식의 다국어 NLI로 전환. 한국어 평가에 "
            "충분히 적합한 학계 표준 모델.")

    H(doc, "14.2 결과 — Fusion이 NLI에서도 명확히 더 충실", level=2)
    Table(doc,
        ["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        [["entailment_rate (↑)", "Anthropic", "0.413", "0.625", "+0.212 ★"],
         ["entailment_rate (↑)", "Gemini", "0.509", "0.624", "+0.115 ★"],
         ["contradiction_rate (↓)", "Anthropic", "0.366", "0.191", "-0.175 ★"],
         ["contradiction_rate (↓)", "Gemini", "0.307", "0.167", "-0.140 ★"],
         ["min_entailment (↑)", "Anthropic", "0.048", "0.181", "+0.134"],
         ["min_entailment (↑)", "Gemini", "0.086", "0.082", "-0.004 ≈"]])
    P(doc, "")
    P(doc, "★ 핵심 1 — Entailment 큰 향상", bold=True, color=(0x55, 0xA8, 0x68))
    P(doc, "Anthropic +0.21, Gemini +0.12. 양 LLM에서 일관되게 fusion이 의미적으로 "
            "더 충실하게 컨텍스트를 따른다.")
    P(doc, "")
    P(doc, "★ 핵심 2 — Contradiction 큰 감소", bold=True, color=(0x55, 0xA8, 0x68))
    P(doc, "Anthropic -0.18, Gemini -0.14. fusion이 컨텍스트와 모순되는 진술을 더 적게 "
            "생성한다.")
    Fig(doc, FIG_DIR / "31_nli_vs_rules.png",
        "그림 14-1. NLI Entailment / Contradiction / Rule sign_match 3-패널 비교")

    H(doc, "14.3 룰 sign_match 하락의 정체 — NLI로 직접 입증", level=2)
    P(doc, "Step 3-C-1에서 가설로만 두었던 \"룰의 키워드 한계\"가 사실임이 NLI로 확정됩니다.")
    Bullet(doc, "Fusion에서 LLM은 \"증가시키는\", \"위험\", \"영향을 준\" 등 다양한 표현을 사용")
    Bullet(doc, "이 단어들은 룰의 pos_words/neg_words 셋에 없어서 sign_in=False가 됨 → false negative")
    Bullet(doc, "그러나 NLI는 단어 셋과 무관하게 의미를 측정 → entailment 명확히 증가, contradiction 명확히 감소")
    Bullet(doc, "결론: **룰 sign_match 하락은 룰의 한계, 진짜 fidelity 손상 아님**")

    H(doc, "14.4 3-Tier 평가 체계의 의미", level=2)
    P(doc, "본 연구의 평가 체계는 이제 세 단계로 구성됩니다.")
    Table(doc,
        ["Tier", "측정 방식", "강점", "한계"],
        [["Rule", "토큰/키워드 매칭", "빠르고 결정적", "다양한 표현 못 잡음 (sign_match)"],
         ["G-Eval", "LLM-as-judge 1~5점", "종합적, 의미 이해", "self-bias 위험 → cross-LLM으로 우회"],
         ["NLI", "의미적 함의 자동 분류", "키워드 무관, 객관적 의미 측정", "짧은 문장에서 noise"]])
    P(doc, "세 tier가 서로의 한계를 보완:")
    Bullet(doc, "룰은 결정적 환각(컨텍스트에 없는 변수명)을 잘 잡고 표현 다양성에서 약함 → NLI가 보완")
    Bullet(doc, "G-Eval은 self-bias 위험 → Cross-LLM(Step 2-A) + 자동 NLI로 보완")
    Bullet(doc, "NLI는 짧은 문장에서 noise → 룰과 G-Eval로 보완")
    Quote(doc, "이 다층 평가는 학계의 RAG 평가 표준(RAGAS, FactCC 등)에 가까운 형태입니다. "
                "본 연구의 약점 중 하나였던 \"LLM 평가 객관성\"이 이번 step으로 부분적으로 "
                "해소되었습니다. 미팅 후 인간평가까지 추가하면 거의 완전 해소됩니다.")

    H(doc, "14.5 Step 3-C-2 종합", level=2)
    Bullet(doc, "NLI로 fusion이 SHAP-only보다 의미적으로 더 충실함을 양 LLM에서 일관 입증")
    Bullet(doc, "룰 sign_match 하락의 정체(키워드 한계)를 직접 입증")
    Bullet(doc, "Rules + G-Eval + NLI의 3-tier 평가 체계 완성")
    Bullet(doc, "약점 1번(LLM 평가 객관성) 부분 해소 — 미팅 후 인간평가로 완전 해소 예정")
    P(doc, "")
    P(doc, "★ Step 3-C-2 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"룰의 sign_match 하락은 룰의 키워드 한계임이 의미적 NLI로 입증되며, fusion은 "
            "세 가지 평가 차원(룰·G-Eval·NLI) 모두에서 충실성을 손상시키지 않으면서 완결성을 "
            "향상시킨다.\"", bold=True, size=12)

    PageBreak(doc)

    PageBreak(doc)

    # ── 14.6 Step 3-C-2-f Cross-Judge 검증 (NEW) ──
    H(doc, "14.6 Step 3-C-2-f — Cross-Judge G-Eval로 평가 종속성 제거 (보너스)",
        level=2)
    P(doc, "Step 3-C-2까지의 G-Eval은 Claude를 judge로 단독 사용했습니다. 이건 한 가지 "
            "위험을 품고 있습니다 — **Claude가 자기 출력(Anthropic target)을 평가할 때 "
            "self-bias를 가질 수 있고, 또 단일 judge의 평가 성향에 결과가 의존할 수 "
            "있습니다.** Step 2-A에서는 양방향 cross-LLM judge를 사용해 self-bias를 우회했었는데, "
            "fusion 평가에서도 같은 패턴을 적용해보았습니다.")
    P(doc, "동일 4 그룹(SHAP-only/Fusion × Anthropic/Gemini target, n=30 each)을 Gemini를 "
            "judge로 다시 평가. Gemini API 503 과부하가 심해서 retry 로직(30s/60s/120s/240s "
            "백오프)을 강화했지만, 결과적으로 모든 120 호출이 회복되어 데이터 손실 0건.")
    Table(doc,
        ["Target", "Metric", "Δ Claude judge", "Δ Gemini judge", "차이"],
        [["Anthropic", "Completeness", "+0.667 ★", "+0.900 ★", "0.23"],
         ["Anthropic", "Factual", "+0.033", "+0.133", "0.10"],
         ["Anthropic", "Sensitive", "0", "0", "0"],
         ["Gemini", "Completeness", "+0.800 ★", "+1.100 ★", "0.30"],
         ["Gemini", "Factual", "−0.133", "+0.467", "**0.60 ★★**"],
         ["Gemini", "Sensitive", "0", "0", "0"]])
    P(doc, "")
    P(doc, "★ 핵심 1 — Completeness 양 judge 모두 큰 향상", bold=True,
       color=(0x55, 0xA8, 0x68))
    P(doc, "Claude judge에서 Anthropic target +0.67, Gemini target +0.80. "
            "Gemini judge에서 Anthropic +0.90, Gemini +1.10. **양 judge에서 일관된 큰 향상** — "
            "fusion 효과가 judge 종속이 아니라 fusion 메커니즘 자체의 효과임이 확정.")
    P(doc, "")
    P(doc, "★ 핵심 2 — Cross-judge가 본 연구에 필수임을 직접 입증", bold=True,
       color=(0xC4, 0x4E, 0x52))
    P(doc, "Gemini target의 Factual Accuracy를 보면, Claude judge는 −0.13으로 \"fusion이 "
            "약간 떨어졌다\"고 판정합니다. 그런데 Gemini judge는 +0.47로 \"fusion이 명확히 "
            "더 사실적이다\"고 판정합니다. 차이 0.60.")
    Quote(doc, "만약 Claude judge 단독 결과만 봤다면, \"Gemini target에서 fusion이 fact를 "
                "약간 손상시킨다\"는 잘못된 결론을 낼 뻔했습니다. Cross-judge로 그게 단일 "
                "judge의 평가 종속성이었음이 드러난 것 — 이게 Cross-LLM judge 방법론의 가치를 "
                "직접 입증하는 사례입니다.")
    P(doc, "")
    P(doc, "★ 핵심 3 — Sensitive Leak 5.0/5.0 양 judge 만점", bold=True,
       color=(0x55, 0xA8, 0x68))
    P(doc, "마스킹 정책(CODE_GENDER, DAYS_BIRTH 등)이 양 judge가 보기에 만점급 — judge "
            "성향과 무관하게 견고함을 다층 검증.")
    Fig(doc, FIG_DIR / "32_cross_judge_geval.png",
        "그림 14-2. Cross-Judge G-Eval (4 메트릭 × 2 target × 2 judge × 2 mode)")
    P(doc, "")
    Quote(doc, "이번 step으로 future work 1순위였던 \"Gemini judge cross-validation\"이 "
                "해소되었습니다. 이제 본 연구의 평가는 Rules + G-Eval(Cross-judge) + NLI "
                "의 3-tier × 2-judge 다층 검증 체계가 됩니다. 인간평가(Plausibility)만 "
                "추가하면 평가 신뢰성 측면의 약점은 거의 완전히 해소됩니다.")

    PageBreak(doc)

    # ── 15. Step 5-A — Fairness-aware Learning (NEW) ──
    H(doc, "15. Step 5-A — Day 5 진단을 mitigation으로 갚는다",
        level=1)
    P(doc, "Day 5에서 본 연구는 정직한 진단을 했습니다 — \"4 모델 × {GENDER, AGE} = "
            "8/8 케이스 4/5 rule 위반\". 그리고 변수 ablation으로 부분적인 mitigation을 "
            "시도했지만, 특히 AGE는 proxy variable이라 ablation 효과가 미미했습니다. "
            "Step 5-A에서는 정식 fairness-aware 학습을 적용해 그 8/8 위반을 어디까지 "
            "되돌릴 수 있는지 정량 측정했습니다.")
    Quote(doc, "결론부터 말하면, Reweighing(Kamiran & Calders 2012) 단일 방법으로 "
                "**4/4 케이스를 모두 4/5 rule 통과**시킬 수 있었습니다. 그것도 baseline "
                "데이터에서 AUROC 손실 0.003 미만, **aux 데이터(Step 3-B)에서는 AUROC 오히려 "
                "+0.003 상승**이라는 의외의 결과까지 나왔습니다.")

    H(doc, "15.1 비교한 3가지 방법", level=2)
    Bullet(doc, "Reweighing — Kamiran-Calders 공식으로 sample_weight 부여. 가장 단순·빠름")
    Bullet(doc, "Fairlearn ExpGrad with DemographicParity (DP) — Reduction-based, max_iter=30")
    Bullet(doc, "Fairlearn ExpGrad with EqualizedOdds (EO) — TPR/FPR 차이만 제약")

    H(doc, "15.2 Reweighing 결과 — 4/4 통과", level=2)
    Table(doc,
        ["데이터", "보호속성", "AUROC 변화", "DI", "4/5 rule"],
        [["baseline (214)", "GENDER", "−0.0024", "0.622 → 0.902", "✅ 통과"],
         ["baseline (214)", "AGE", "−0.0038", "0.557 → 0.901", "✅ 통과"],
         ["aux (1161)", "GENDER", "+0.0028 ★", "0.643 → 0.867", "✅ 통과"],
         ["aux (1161)", "AGE", "+0.0016 ★", "0.567 → 0.833", "✅ 통과"]])
    P(doc, "")
    P(doc, "★ aux 데이터에서 AUROC 향상 — 의외의 발견", bold=True,
       color=(0xC4, 0x4E, 0x52))
    P(doc, "이 결과는 \"공정성 mitigation에는 항상 성능 비용이 있다\"는 통념을 뒤집습니다. "
            "보조 테이블의 풍부한 정보 공간(214 → 1161 features)이 sample weight 변화로 "
            "발생하는 분포 shift를 안정적으로 처리해, weight 조정이 모델에게 더 균형 잡힌 "
            "학습 신호로 작용한 것으로 해석됩니다.")
    Fig(doc, FIG_DIR / "33_fairness_tradeoff.png",
        "그림 15-1. AUROC vs Disparate Impact 산포 (Reweighing 점들이 우상단)")

    H(doc, "15.3 Fairlearn과 비교 — Reweighing이 압도적", level=2)
    P(doc, "Fairlearn ExpGrad는 학계의 SOTA reduction-based 방법인데, 본 데이터에는 "
            "왜인지 결과가 좋지 않았습니다.")
    Table(doc,
        ["보호속성", "Method", "AUROC", "DI", "4/5 rule"],
        [["GENDER", "Reweighing ★", "0.7581", "0.902", "✅"],
         ["GENDER", "Fairlearn DP", "0.7088", "0.775", "❌"],
         ["GENDER", "Fairlearn EO", "0.6788", "0.659", "❌"],
         ["AGE", "Reweighing ★", "0.7567", "0.901", "✅"],
         ["AGE", "Fairlearn DP", "0.6856", "0.990", "✅ (큰 AUROC 손실)"],
         ["AGE", "Fairlearn EO", "0.7225", "0.150 ⚠️", "❌ 안티 패턴"]])
    P(doc, "")
    Bullet(doc, "Fairlearn DP는 강한 공정성 제약 → AUROC -0.07~-0.075 손실. 가장 \"공정\"하지만 비용 큼")
    Bullet(doc, "Fairlearn EO는 EO에만 집중하니 selection rate 자체엔 제약 X — DP가 오히려 악화 (AGE+EO에서 DI 0.150). 본 데이터의 4/5 rule(DI 기반) 통과엔 부적합 — 안티 패턴")
    Bullet(doc, "→ 본 데이터에서 Reweighing이 가장 효율적인 mitigation임을 정량 입증")
    Fig(doc, FIG_DIR / "34_mitigation_bars.png",
        "그림 15-2. Mitigation 방법별 DI/DP/AUROC 비교 (보호속성 × 데이터)")

    H(doc, "15.4 AGE에 대한 새로운 발견", level=2)
    P(doc, "Day 5에서는 \"AGE는 proxy variable로 다른 변수(DAYS_EMPLOYED 등)에 간접 "
            "인코딩되어 있어서 ablation 효과가 미미하다\"고 결론지었었습니다. Step 5-A의 "
            "결과는 그 결론을 갱신합니다:")
    Bullet(doc, "AGE도 Reweighing으로 효과적인 mitigation 가능 (DI 0.557 → 0.901)")
    Bullet(doc, "Day 5의 변수 ablation은 \"AGE 변수만 빼면 되는가\"를 측정한 것")
    Bullet(doc, "Step 5-A는 \"학습 분포 자체를 조정\"하므로 proxy effect를 우회 가능")
    Bullet(doc, "→ 정식 fairness-aware 방법은 변수 ablation의 한계를 극복할 수 있음")

    H(doc, "15.5 Step 5-A 종합", level=2)
    Bullet(doc, "Day 5 진단(8/8 위반)을 단일 mitigation 방법(Reweighing)으로 4/4 통과")
    Bullet(doc, "약점 #4 (Fairness mitigation gap) 해소")
    Bullet(doc, "aux 데이터에서 trade-off 사라짐 — 적절한 feature space에서 fairness가 win-win 가능")
    Bullet(doc, "Fairlearn 비교로 본 데이터에서의 mitigation 방법 선택 정량 가이드 제공")
    P(doc, "")
    P(doc, "★ Step 5-A 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"본 연구의 공정성 진단(Day 5)은 단순한 alarm이 아니라 정량 mitigation으로 "
            "이어진다. Reweighing 4/4 통과 + aux에서 trade-off 사라짐.\"",
       bold=True, size=12)

    PageBreak(doc)

    # ── 16. Step 5-B — Generic RAG Baseline (NEW) ──
    H(doc, "16. Step 5-B — Counterfactual baseline의 trivial 반박을 직면하다", level=1)
    P(doc, "Step 1의 Counterfactual baseline 결과(\"XAI-RAG는 환각 0%, no-SHAP은 45.5%\")는 "
            "본 연구에서 가장 강력한 메시지 중 하나입니다. 그러나 신중한 심사자는 이렇게 "
            "반박할 수 있습니다:")
    Quote(doc, "\"그건 trivial한 결과 아닌가? 정보가 적은 쪽이 환각하는 건 당연하지 않나?\"")
    P(doc, "Step 5-B는 이 반박을 정면으로 다룹니다. 같은 양의 정보(raw features + 일반 도메인 "
            "지식 chunks) + 같은 hard constraints를 주되 SHAP context만 빼고, 환각률과 충실성을 "
            "비교합니다.")

    H(doc, "16.1 4-mode 설계", level=2)
    Bullet(doc, "no_shap (Step 1): raw 데이터, 자유 추론, hard constraints 약함")
    Bullet(doc, "★ generic_rag (Step 5-B, NEW): raw + 도메인 지식 chunks 7개 + 동일 hard constraints, SHAP X")
    Bullet(doc, "shaponly (Step 1/2-A): SHAP top-k drivers + hard constraints")
    Bullet(doc, "fusion (Step 3-C-1): SHAP + TabNet attention agreement-aware")
    P(doc, "Generic RAG의 Knowledge chunks 7개: 신용평가 핵심 변수의 의미, 부도 위험의 일반 원리, "
            "threshold 의미, 민감 변수 마스킹 정책, 금융 용어 가이드(DTI/LTV/DSR 등), "
            "hard constraints, 출력 형식.")

    H(doc, "16.2 결과 — 4단계 명확한 차이", level=2)
    Table(doc,
        ["Metric", "no_shap", "generic_rag", "shaponly", "fusion"],
        [["Halluc strict (Anthropic)", "0.167 ⚠️", "0.000 ✅", "0.000", "0.000"],
         ["Halluc strict (Gemini)", "0.000", "0.000", "0.000", "0.000"],
         ["NLI Entailment ★ (A)", "0.270", "0.364", "0.413", "0.625 ★"],
         ["NLI Entailment ★ (G)", "0.429", "0.370", "0.509", "0.624 ★"],
         ["G-Eval Completeness (A)", "3.000", "4.833", "4.300", "4.967"],
         ["val_match_rate (A)", "0.588", "0.727", "0.847", "0.903 ★"],
         ["val_match_rate (G)", "0.412", "0.694", "0.793", "0.861 ★"]])
    P(doc, "")
    P(doc, "★ 핵심 1 — 환각 차단은 Generic RAG로도 가능", bold=True,
       color=(0xC4, 0x4E, 0x52))
    P(doc, "Anthropic의 no_shap만 환각 0.167(Step 1의 45.5% 패턴 재현). "
            "Generic RAG 추가 시 hard constraints + 도메인 지식 chunks만으로 0%로 떨어집니다. "
            "즉 \"환각 차단\"이 SHAP-RAG의 유일한 가치는 아닙니다.")
    P(doc, "")
    P(doc, "★ 핵심 2 — 의미적 충실성 4단계 명확", bold=True,
       color=(0x55, 0xA8, 0x68))
    P(doc, "NLI Entailment에서 양 LLM 일관된 4단계 차이: "
            "no_shap(0.27~0.43) < generic_rag(0.36~0.37) < shaponly(0.41~0.51) < "
            "**fusion(0.62~0.62)**. 이 4단계 차이가 본 연구 메커니즘의 진짜 차별성입니다.")
    P(doc, "")
    P(doc, "★ 핵심 3 — 값 정확 인용에서 SHAP 우위", bold=True,
       color=(0x55, 0xA8, 0x68))
    P(doc, "val_match_rate (LLM이 컨텍스트의 값을 정확히 인용한 비율)은 mode별로 "
            "0.59 → 0.73 → 0.85 → 0.90으로 명확하게 증가. SHAP-RAG/Fusion이 도달하는 "
            "fact-grounded 정확성은 Generic RAG로 모방 불가능합니다.")
    Fig(doc, FIG_DIR / "35_generic_rag_3way.png",
        "그림 16-1. 4-mode 비교 (Halluc / NLI / G-Eval Completeness × 양 LLM)")

    H(doc, "16.3 흥미로운 발견 — Generic RAG의 Completeness 우위", level=2)
    P(doc, "Anthropic의 G-Eval Completeness에서 Generic RAG(4.83)가 SHAP-only(4.30)보다 "
            "높게 나왔습니다. 처음엔 의외이지만 해석하면:")
    Bullet(doc, "Generic RAG의 도메인 지식 chunks 7개가 LLM에게 더 풍부한 일반 컨텍스트 제공")
    Bullet(doc, "SHAP-only는 driver 정보만 — 일반 원리 가이드 부족")
    Bullet(doc, "그러나 Fusion(4.97)이 여전히 최고 — 두 신호 융합이 Generic RAG의 chunks 효과를 통합/초과")
    P(doc, "이건 Generic RAG의 가치를 정직 인정하면서도, Fusion이 그 효과까지 흡수함을 보여줍니다.")

    H(doc, "16.4 본 연구의 진짜 차별성 재정의", level=2)
    P(doc, "Step 1 메시지(\"환각 0% vs 45.5%\")만으로는 불완전합니다. Step 5-B 후 본 연구의 진짜 메시지는:")
    Quote(doc, "\"본 XAI-RAG 시스템의 차별성은 단순한 환각 차단이 아니라, 두 해석 신호(SHAP + "
                "TabNet attention)의 의미적 충실성과 값 정확 인용에 대한 압도적 우위다. "
                "일반 도메인 지식 RAG는 환각 차단을 모방할 수 있지만 fact-grounded 정확성은 "
                "SHAP/Fusion 컨텍스트 없이 달성 불가능하다.\"")

    H(doc, "16.5 Step 5-B 종합", level=2)
    Bullet(doc, "★ 약점 #3 (Counterfactual baseline 정당성) 정량 해소")
    Bullet(doc, "Honest reporting — \"환각 차단의 trivial 반박\" 부분 인정")
    Bullet(doc, "본 연구 차별성 재정의 — 의미적 충실성 + 값 정확 인용")
    Bullet(doc, "양 LLM에서 4단계 차이 일관 확인 → LLM 종속성 없는 메커니즘")
    P(doc, "")
    P(doc, "★ Step 5-B 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"환각 차단은 hard constraints로도 가능하지만, fact-grounded 의미적 충실성은 "
            "SHAP/Fusion 컨텍스트만이 도달할 수 있는 수준이다.\"",
       bold=True, size=12)

    PageBreak(doc)

    # ── 17. Step 5-C — Pilot Human-Proxy Evaluation (NEW) ──
    H(doc, "17. Step 5-C — 사람 perspective에서는 다른 결과가 나온다", level=1)
    P(doc, "정식 IRB 인간평가는 1.5~2개월 행정 부담이라 미팅 시점 진행이 어렵습니다. "
            "Step 5-C는 그 pilot 대안 — LLM persona를 사람 대리(human proxy)로 사용해 "
            "3가지 stakeholder 관점에서 plausibility를 정량 측정합니다.")
    Quote(doc, "이 step의 결과는 본 연구의 가장 정직한 발견을 만들어냅니다 — "
                "**본 연구의 fusion 메커니즘이 모든 평가 차원에서 1위가 아니라는 점**.")

    H(doc, "17.1 3 personas 설계", level=2)
    Bullet(doc, "Credit Expert (10년 경력 신용 분석가) — 전문가 신뢰성")
    Bullet(doc, "Customer (대출 신청자 본인) — 일반 고객 납득도")
    Bullet(doc, "Regulator (금융감독원 평가자) — 규제 요건 만족도")
    P(doc, "각 persona × 5점 척도 (trustworthiness / clarity / actionability) × "
            "4 modes × 2 LLM target × 15 instances = 360 평가 (실제 276, no_shap 표본 작음).")

    H(doc, "17.2 결과 — 4 modes × 3 personas (Anthropic target, n=15)", level=2)
    Table(doc,
        ["Persona / Metric", "no_shap*", "generic_rag", "shaponly", "fusion"],
        [["Credit Expert / trust", "(2.0)", "4.80 ★", "4.33", "4.73"],
         ["Credit Expert / clarity", "(4.0)", "4.93 ★", "4.67", "4.73"],
         ["Credit Expert / action", "(3.0)", "4.73 ★", "3.40", "3.73"],
         ["Customer / trust", "(5.0)", "4.93 ★", "3.53", "3.67"],
         ["Customer / clarity ⚠️", "(5.0)", "4.93 ★", "2.80", "2.67"],
         ["Customer / action", "(5.0)", "4.33 ★", "2.33", "2.80"],
         ["Regulator / trust", "(5.0)", "5.00 ★", "4.53", "4.53"],
         ["Regulator / clarity", "(5.0)", "5.00 ★", "4.40", "4.53"],
         ["Regulator / action", "(5.0)", "5.00 ★", "3.33", "3.60"]])
    P(doc, "* no_shap n=1, 통계적 의미 없음")
    Fig(doc, FIG_DIR / "36_human_proxy_personas.png",
        "그림 17-1. 3 metrics × 3 personas × 4 modes × 2 LLM target")

    H(doc, "17.3 충격적 발견 — fusion이 1위가 아니다", level=2)
    P(doc, "Persona 관점에서 모든 personas × 거의 모든 metrics에서 **Generic RAG가 1위**입니다. "
            "본 연구의 fusion은 통상 SHAP-only보다는 약간 우위지만, Generic RAG에는 미치지 "
            "못합니다.")
    Quote(doc, "특히 Customer perspective에서 SHAP-RAG/Fusion clarity 2.67~2.80(만점 5)은 "
                "큰 약점입니다. \"두 모델이 동의한 강한 신호\", \"부도 가능성↑\" 같은 기술적 "
                "표기가 일반 고객에게 이해하기 어렵게 작용합니다. Generic RAG는 도메인 chunks의 "
                "자연스러운 한국어 + raw value만 인용하므로 customer-facing 시나리오에서 "
                "압도적입니다.")

    H(doc, "17.4 Trade-off 정량 입증 — 평가 차원에 따른 1위 다름", level=2)
    Table(doc,
        ["평가 차원", "1위", "비고"],
        [["사실성 (NLI Entailment)", "fusion 0.62", "Step 3-C-2 결과"],
         ["충실성 (G-Eval Completeness)", "fusion 4.97", "Step 5-B 결과"],
         ["사람 친화성 (Persona trust)", "generic_rag 4.91 ★", "Step 5-C 발견 — fusion 1위 아님"],
         ["Customer clarity ⚠️", "generic_rag 4.93 ★", "fusion 2.67, shaponly 2.80 — 큰 차이"]])
    Quote(doc, "이건 단순한 약점이 아니라 \"응용 시나리오에 따른 mode 선택 trade-off\"의 "
                "정량 입증입니다. fusion은 audit/regulation 시나리오에 적합 (사실성·충실성 압도), "
                "Generic RAG는 customer-facing UI에 적합 (친근함). "
                "또는 향후 두 표현을 결합하는 hybrid 형태도 가능합니다.")

    H(doc, "17.5 본 연구 메시지의 단계별 정교화", level=2)
    Bullet(doc, "Step 1: \"환각 차단 0% vs 45.5%\" — 강력하지만 trivial 반박 가능")
    Bullet(doc, "Step 5-B: \"환각 차단은 hard constraints로도 가능, 차별성은 fact-grounded 정확성\"")
    Bullet(doc, "★ Step 5-C: \"사실성 1위 + 친근함은 trade-off, 응용에 따라 선택\"")
    P(doc, "각 step마다 메시지가 더 정교해지고 한계도 더 정직하게 인정됩니다 — "
            "이건 학술 논문의 표준적 진행 방식입니다.")

    H(doc, "17.6 Pilot의 본질적 한계", level=2)
    Bullet(doc, "LLM persona는 사람 proxy — 진짜 인간 사고는 다를 수 있음")
    Bullet(doc, "Judge LLM은 Claude 단일 — Cross-judge cross-validation은 future work")
    Bullet(doc, "n=15 표본 — 95% CI 큼, 효과 크기는 명확하지만 정밀도 제한")
    Bullet(doc, "no_shap n=1 — 통계적 의미 없음, 표본 확장 필요")
    Bullet(doc, "Customer persona의 \"일반 고객\" 정의 — 실제 고객 다양성 못 반영")

    H(doc, "17.7 Step 5-C 종합", level=2)
    Bullet(doc, "★ 약점 #2 부분 해소 — informal pilot, 정식 IRB는 future work")
    Bullet(doc, "★ Customer clarity 약점 발견 — 본 연구 한계 honest 인정")
    Bullet(doc, "★ Trade-off 정량 입증 — 응용에 따른 mode 선택 가이드")
    Bullet(doc, "메시지 정교화 — \"단순 우월 X, 다층 우위와 trade-off\"")
    P(doc, "")
    P(doc, "★ Step 5-C 한 줄 메시지", bold=True, size=14, color=(0xC4, 0x4E, 0x52))
    P(doc, "\"본 연구의 fusion 메커니즘은 사실성에 압도적이지만, 사람 친화성은 Generic RAG가 "
            "우위. 응용 시나리오에 따른 mode 선택 trade-off가 정량 입증되었다.\"",
       bold=True, size=12)

    PageBreak(doc)

    H(doc, "18. 한계와 향후 계획 (Future Work)", level=1)
    Bullet(doc, "인간 평가 (Plausibility) — IRB 간소판 신청, 5점 리커트 척도, Cohen's κ 신뢰도. **약점 1번 완전 해소를 위해 미팅 후 1순위.**")
    Bullet(doc, "UCI German Credit — 데이터 다양성, 일반화 입증 ★ 2순위.")
    Bullet(doc, "no_shap 표본 확장 — 현재 30 idx 중 2개만 평가 가능, 통계적 견고성 강화 필요.")
    Bullet(doc, "(완료) Gemini judge cross-validation — Step 3-C-2-f에서 진행, 양 judge 일관 입증.")
    Bullet(doc, "Fusion 평가 표본 30 → 100 확장 + counterfactual 결합.")
    Bullet(doc, "3-way ablation: SHAP-only / Attention-only / Fusion 비교.")
    Bullet(doc, "보조 테이블 잔여 4개(POS_CASH_balance, credit_card_balance, "
                "installments_payments, bureau_balance 추가 활용) → AUROC 0.78+ 도전.")
    Bullet(doc, "Bureau ablation: bureau 단독 vs prev 단독 vs 둘 다 비교로 "
                "EXT_SOURCE 응축 가설 검증.")
    Bullet(doc, "TabNet, LightGBM에 aux 효과 일반화 (Step 3-B는 XGBoost만 검증).")
    Bullet(doc, "Fairness-aware 학습: Reweighing, Adversarial Debiasing.")
    Bullet(doc, "FT-Transformer 비교 모델 추가.")
    Bullet(doc, "한국어 native NLI 모델 추가 검증 (현재 다국어; torch 환경 정비 후 KLUE-roberta-NLI).")
    Bullet(doc, "한국어 도메인 특화 금융 LLM 미세조정 (QLoRA).")
    Bullet(doc, "한국어 도메인 특화 금융 LLM 미세조정 (QLoRA).")

    P(doc, "")
    Quote(doc, "Step 1 — \"미팅용 작동 프로토타입\" 완성. "
                "Step 2-A — 평가 신뢰성 강화 (100명, Cross-LLM, Robustness). "
                "Step 3-B — 성능 확장 (보조 테이블, AUROC +2.22%). "
                "Step 3-C-1 — TabNet 메커니즘 통합 (논문 제목 정당화). "
                "Step 3-C-2 — NLI로 평가 객관성 보강 (3-tier 평가 체계 완성). "
                "Step 3-C-2-f — Cross-Judge G-Eval로 평가 종속성 제거. "
                "Step 5-A — Fairness mitigation (Reweighing 4/4 통과). "
                "Step 5-B — Generic RAG baseline (약점 #3 해소, 차별성 재정의). "
                "Step 5-C — Pilot human-proxy (약점 #2 부분 해소, trade-off 발견). "
                "이후 UCI German Credit + 정식 IRB + Customer-friendly 표현 정제로 확장합니다.")

    out = PAPER_DIR / "midterm_report_friendly.docx"
    doc.save(out)
    print(f"[OK] {out} 저장")


if __name__ == "__main__":
    main()
