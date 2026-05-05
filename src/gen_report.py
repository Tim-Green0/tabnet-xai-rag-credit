"""미팅용 docx 보고서 생성기.

흐름:
  1. 표지 + 1줄 요약
  2. 연구 목적·기여 (계획서 1.2 기반)
  3. 데이터·전처리
  4. 모델 비교 (5-fold CV)
  5. SHAP × 어텐션 일관성 (RQ2)
  6. 공정성 진단 + Mitigation
  7. XAI-RAG 자연어 설명 (Demo 인스턴스)
  8. 정량 평가 (RQ3, 환각 0%)
  9. 향후 계획 (Future Work)

실행:  PYTHONIOENCODING=utf-8 .venv/Scripts/python.exe -m src.gen_report
산출:  paper/midterm_report.docx
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

from src.utils import RESULTS_DIR

PAPER_DIR = Path("paper")
PAPER_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR = Path("figures")


# ─────────────────────────────────────────────────────────────
# 헬퍼
# ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, color: str):
    """테이블 셀 배경색."""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(16 if level == 1 else 13)
    return h


def add_para(doc, text: str, bold: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.bold = bold
    # East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    from docx.oxml import OxmlElement
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rPr.append(rFonts)
    return p


def add_table_from_data(doc, headers: list, rows: list, header_bg: str = "4C72B0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_bg)
    # Rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Malgun Gothic"
            run.font.size = Pt(10)
    return table


def add_figure(doc, path: Path, caption: str | None = None, width_cm: float = 15):
    if not path.exists():
        add_para(doc, f"[그림 누락: {path}]")
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def load_json(p: Path) -> dict:
    return json.loads(Path(p).read_text(encoding="utf-8"))


# ─────────────────────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # ── 표지 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("석사 논문 중간 보고")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("정형 데이터 특화 딥러닝(TabNet)과 LLM 기반 XAI-RAG를\n"
                       "활용한 설명 가능한 신용 평가 및 사용자 맞춤형 리포트 생성")
    run.font.size = Pt(13)
    run.font.name = "Malgun Gothic"

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("\n").font.size = Pt(10)
    for txt in ["A70067 오현택", "지도교수: 박운상", "보고일: 2026-05-10 (예정)"]:
        r = info.add_run(txt + "\n")
        r.font.name = "Malgun Gothic"
        r.font.size = Pt(11)

    doc.add_paragraph()

    # 1줄 메시지
    msg = doc.add_paragraph()
    msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = msg.add_run("[핵심 메시지] 본 연구의 XAI-RAG는 두 상용 LLM 모두에서 "
                     "Hallucination Rate 0.000을 달성 — 환각 차단 효과의 LLM 종속성 없음을 입증")
    r.font.bold = True
    r.font.name = "Malgun Gothic"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC4, 0x4E, 0x52)

    doc.add_page_break()

    # ── 1. 연구 목적·기여 ──
    add_heading(doc, "1. 연구 목적과 기여", level=1)
    add_para(doc, "정형 데이터 신용 평가에서 (1) 고성능 예측, (2) 투명한 수학적 근거, "
                  "(3) 사용자 친화적 자연어 설명 세 요소를 일관된 파이프라인으로 통합하고, "
                  "그 결과의 신뢰성을 정량 검증한다.")
    add_heading(doc, "1.1 핵심 차별점", level=2)
    add_bullet(doc, "TabNet 어텐션 마스크와 SHAP의 일관성을 정량 분석 (대부분 연구가 사후 SHAP만 사용)")
    add_bullet(doc, "SHAP 결과를 'retrieved evidence'로 정의해 LLM 환각을 원천 차단하는 XAI-RAG 구조")
    add_bullet(doc, "Faithfulness · Hallucination · G-Eval 다층 정량 평가 프로토콜")
    add_bullet(doc, "성별·연령 4종 공정성 지표 + ablation 기반 mitigation 비교")

    # ── 2. 데이터 ──
    add_heading(doc, "2. 데이터 및 전처리", level=1)
    add_para(doc, "데이터셋: Kaggle Home Credit Default Risk (메인 테이블), 307,511행 × 122컬럼.")
    add_para(doc, "TARGET 분포 8.07% (불균형) → AUROC, AUPRC, KS 위주 평가 + class_weight 균형.")
    add_bullet(doc, "전처리: 결측 50%+ 컬럼은 *_MISSING_FLAG 추가 후 보존")
    add_bullet(doc, "범주형: cardinality ≤8 → one-hot, OCCUPATION/ORGANIZATION → target encoding")
    add_bullet(doc, "DAYS_EMPLOYED sentinel 365243(≈18%)는 NaN+EMPLOYED_FLAG 처리")
    add_bullet(doc, "RobustScaler 적용 (TabNet/Logistic용), 60/20/20 stratified split, SEED=42")
    add_bullet(doc, "최종 feature 수: 122 → 214")

    # ── 3. 모델 성능 ──
    add_heading(doc, "3. 모델 비교 (5-fold Stratified CV, test set)", level=1)
    add_para(doc, "단발 학습이 아닌 5-fold mean ± std로 안정성 검증.")
    add_table_from_data(doc,
        headers=["모델", "AUROC", "AUPRC", "KS", "F1", "비고"],
        rows=[
            ["XGBoost", "0.7587 ± 0.0008", "0.2445 ± 0.0011",
             "0.3846 ± 0.0015", "0.2698 ± 0.0047", "1등 (5/5 fold)"],
            ["LightGBM", "0.7544 ± 0.0009", "0.2402 ± 0.0018",
             "0.3788 ± 0.0028", "0.2584 ± 0.0052", "—"],
            ["Logistic", "0.7544 ± 0.0001", "0.2343 ± 0.0006",
             "0.3804 ± 0.0010", "0.2631 ± 0.0087", "선형 모델 안정성 최고"],
            ["TabNet", "0.7518 ± 0.0017", "0.2331 ± 0.0023",
             "0.3749 ± 0.0056", "0.2657 ± 0.0079", "어텐션 해석성 제공"],
        ])
    add_para(doc, "")
    add_figure(doc, FIG_DIR / "13_cv_comparison.png",
                "그림 1. 5-fold CV 비교 (test, mean ± std)")

    # ── 4. SHAP × Attention ──
    add_heading(doc, "4. SHAP × 어텐션 일관성 (RQ2)", level=1)
    add_para(doc, "TabNet 어텐션 마스크와 SHAP global importance 간의 정량 일관성 분석.")
    add_table_from_data(doc,
        headers=["지표", "값", "해석"],
        rows=[
            ["Spearman ρ (전체 214 변수)", "0.117 (p=0.089)", "약한 양의 상관"],
            ["Spearman ρ (Top 50 합집합)", "−0.195", "음수 — 미세 영역에서 상보적"],
            ["Top-20 교집합", "9 / 20 (Jaccard 0.29)", "핵심 변수는 일치"],
        ])
    add_para(doc, "")
    add_para(doc, "교집합 9개: EXT_SOURCE_2, EXT_SOURCE_3, DAYS_EMPLOYED, "
                  "ORGANIZATION_TYPE, NAME_CONTRACT_TYPE_Revolving loans, "
                  "CODE_GENDER_M 등.", size=10)
    add_para(doc, "→ 두 방법은 핵심 변수에서 일관적이지만 미세 순위에서는 상보적. "
                  "본 연구의 다층 해석 정당성을 데이터로 뒷받침.", bold=True)
    add_figure(doc, FIG_DIR / "16_attention_vs_shap_scatter.png",
                "그림 2. 어텐션 vs SHAP 일관성 산점도")

    # ── 5. 공정성 ──
    add_heading(doc, "5. 공정성 진단 + Mitigation", level=1)
    add_para(doc, "4개 모델 × 보호속성 {GENDER, AGE} = 8건 모두 4/5 rule 위반 (DI < 0.8).")
    add_para(doc, "보호 속성 컬럼 제거 후 재학습한 ablation 결과:")
    add_table_from_data(doc,
        headers=["모델 × 속성", "AUROC Δ", "DP Δ", "DI Δ"],
        rows=[
            ["XGBoost × GENDER", "−0.005", "0.164 → 0.105 (−36%)", "0.622 → 0.718"],
            ["TabNet × GENDER", "−0.010", "0.156 → 0.093 (−40%)", "0.621 → 0.757"],
            ["XGBoost × AGE", "−0.005", "0.197 → 0.179 (−9%)", "0.498 → 0.507"],
            ["TabNet × AGE", "−0.010", "0.185 → 0.188 (+2%)", "0.504 → 0.513"],
        ])
    add_para(doc, "")
    add_bullet(doc, "GENDER ablation은 효과적 (DP 30~40% 감소, AUROC 1% 미만 손실)")
    add_bullet(doc, "AGE ablation은 효과 미미 — 연령이 DAYS_EMPLOYED·DAYS_REGISTRATION 등 "
                    "다른 변수에 간접 인코딩된 proxy variable 문제")
    add_bullet(doc, "4/5 rule은 여전히 미통과 → 본격적 fairness-aware 학습 필요 (future work)")
    add_figure(doc, FIG_DIR / "19_fairness_mitigation.png",
                "그림 3. baseline vs ablated 공정성 비교")

    # ── 6. XAI-RAG 데모 ──
    add_heading(doc, "6. XAI-RAG 자연어 설명 (Demo: idx 59291)", level=1)
    walk = load_json(RESULTS_DIR / "demo_walkthrough.json")
    s = walk["sample"]
    p = walk["prediction"]
    add_para(doc, f"인스턴스: 실제 정답 = {'부도(1)' if s['true_label'] == 1 else '정상(0)'}, "
                  f"P(default) = {p['default_proba']:.4f}, 결정 = {p['decision']} "
                  f"(True Positive)")
    add_heading(doc, "6.1 SHAP 기반 거절 측 Top 5", level=2)
    rows = [[d["rank"], d["feature"], d["value"], f"{d['shap']:+.4f}"]
            for d in walk["context"]["top_drivers_for_default"]]
    add_table_from_data(doc, ["rank", "feature", "value", "SHAP"], rows)
    add_para(doc, "")

    add_heading(doc, "6.2 두 LLM 자연어 설명 비교", level=2)
    for prov_key, prov_label in [("gemini", "Gemini 2.5 Flash"),
                                    ("anthropic", "Claude Sonnet 4.5")]:
        if prov_key in walk["llm_outputs"] and "explanation" in walk["llm_outputs"][prov_key]:
            ll = walk["llm_outputs"][prov_key]
            add_para(doc, f"[{prov_label}] elapsed={ll['elapsed_sec']:.1f}s, "
                          f"tokens={ll.get('total_tokens', '?')}", bold=True, size=10)
            for line in ll["explanation"].split("\n"):
                if line.strip():
                    add_para(doc, line.strip(), size=10)
            add_para(doc, "")

    # ── 7. 평가 ──
    add_heading(doc, "7. 정량 평가 (RQ3 — 환각 차단 효과)", level=1)
    cmp_df = pd.read_csv(RESULTS_DIR / "llm_comparison.csv")
    add_para(doc, "평가 차원: Faithfulness, Hallucination, G-Eval(Gemini self-judge), 효율성. "
                  "10 샘플(REJECT 5 + APPROVE 5)에 대해 측정.", size=10)
    add_table_from_data(doc,
        headers=["지표", "Gemini 2.5 Flash", "Claude Sonnet 4.5", "비고"],
        rows=[
            ["Hallucination Rate (strict)", "0.000 ± 0.000", "0.000 ± 0.000",
             "★ XAI-RAG 환각 차단 입증"],
            ["Hallucination Rate (broad)", "0.000 ± 0.000", "0.000 ± 0.000",
             "컨텍스트 외 변수 0건"],
            ["val_match_rate", "0.811 ± 0.123", "0.901 ± 0.111", "Claude 우위"],
            ["sign_match_rate", "0.783 ± 0.209", "0.867 ± 0.233", "Claude 우위"],
            ["G-Eval factual_accuracy", "5.0 ± 0.0 / 5", "(skip)", "self-judge 만점"],
            ["G-Eval sensitive_leak", "5.0 ± 0.0 / 5", "(skip)", "민감변수 마스킹 완벽"],
            ["G-Eval style", "5.0 ± 0.0 / 5", "(skip)", "고객 친화적"],
            ["elapsed (sec, per call)", "12.7 ± 4.5", "8.4 ± 0.8", "Claude 빠름"],
            ["total tokens (per call)", "4,155 ± 834", "2,500 ± 83", "Claude 효율적"],
        ])
    add_para(doc, "")
    add_figure(doc, FIG_DIR / "21_llm_comparison.png",
                "그림 4. Gemini vs Claude — 룰 기반·G-Eval·효율성")

    add_heading(doc, "7.1 핵심 메시지", level=2)
    add_bullet(doc, "두 상용 LLM 모두에서 Hallucination Rate 0.000 — XAI-RAG 환각 차단 효과의 "
                    "LLM 종속성 없음")
    add_bullet(doc, "G-Eval factual_accuracy / sensitive_leak / style 모두 만점 (5.0/5)")
    add_bullet(doc, "Faithfulness 룰 기반 정합도 약간의 차이는 출력 스타일(픽셀 일치 vs 자연 반올림)")
    add_bullet(doc, "Claude가 효율 우위 — 시간 −34%, 토큰 −40%")

    # ── 8. Counterfactual Baseline ──
    add_heading(doc, "8. Counterfactual Baseline — XAI-RAG vs no-SHAP", level=1)
    add_para(doc, "RQ3 강한 검증: 동일 11샘플에 대해 SHAP 컨텍스트 없이 raw 데이터만 "
                  "LLM에 주는 baseline 실험. 환각률 비교.")
    cmp_path = RESULTS_DIR / "baseline_comparison.csv"
    if cmp_path.exists():
        cmp = pd.read_csv(cmp_path)
        rows = [
            [r["provider"].title(),
             f"{r['xai_rag_halluc_strict_mean']:.3f} ± {r['xai_rag_halluc_strict_std']:.3f}",
             f"{r['baseline_halluc_strict_mean']:.3f} ± {r['baseline_halluc_strict_std']:.3f}",
             f"{int(r['baseline_outside_dataset_total'])}/{int(r['baseline_n_candidates_total'])}"]
            for _, r in cmp.iterrows()
        ]
        add_table_from_data(doc,
            ["LLM", "XAI-RAG halluc", "Baseline halluc", "환각/총 후보 (baseline)"],
            rows)
    add_para(doc, "")
    add_para(doc,
        "Claude Sonnet 4.5의 baseline에서 환각률 45.5% (XAI-RAG는 0%) — 결정적 차이.",
        bold=True)
    add_para(doc, "Claude baseline 환각 사례:", bold=True, size=10)
    add_bullet(doc, "DTI, LTV, DSR — Home Credit 데이터셋에 없는 일반 금융 비율 약어를 "
                    "LLM이 자체 학습된 도메인 지식으로 끌어와 사용")
    add_bullet(doc, "'햇살론, 미소금융' — 데이터에 없는 특정 금융 상품명을 권고에 포함")
    add_bullet(doc, "'1588-XXXX' 같은 가짜 고객센터 번호 생성")
    add_bullet(doc, "DEF_30_CNT 등 변수명 잘림 — 데이터의 실제 변수명 부정확 인용")
    add_para(doc, "")
    add_para(doc,
        "Gemini의 baseline 측정값 0%는 환각이 없는 것이 아니라, "
        "한국어 자연어로 의역해 영문 변수명 정규식 룰로 잡히지 않은 것. "
        "측정 한계 — Cross-LLM judge 등 의미 단위 평가가 future work.", size=10)
    add_figure(doc, FIG_DIR / "23_baseline_vs_xairag.png",
                "그림 5. XAI-RAG vs baseline (no SHAP) — 환각률 비교")

    # ── 9. Step 2-A 평가 신뢰성 강화 (NEW) ──
    add_heading(doc, "9. Step 2-A — 평가 신뢰성 강화 (100명 표본)", level=1)
    add_para(doc, "Step 1의 가장 큰 약점이었던 평가 표본 10명을 100명으로 확장하고, "
                  "Cross-LLM G-Eval, Counterfactual Test, Robustness 평가를 추가했다.")

    add_heading(doc, "9.1 100명 표본 환각률 — 0% 유지", level=2)
    add_table_from_data(doc,
        headers=["LLM", "Hallucination Rate (strict)", "n", "비고"],
        rows=[
            ["Gemini 2.5 Flash", "0.000 ± 0.000", "100", "Step 1 (n=10) 결과 견고"],
            ["Claude Sonnet 4.5", "0.000 ± 0.000", "100", "Step 1 (n=10) 결과 견고"],
        ])
    add_para(doc, "")
    add_para(doc, "10명 → 100명 확장 후에도 환각률 0% 유지. 통계적 안정성 확보.",
              bold=True)

    add_heading(doc, "9.2 Cross-LLM G-Eval (self-bias 우회)", level=2)
    add_table_from_data(doc,
        headers=["Judge → Target", "Factual", "Complete", "Sensitive", "Style"],
        rows=[
            ["Gemini → Gemini (Step 1, n=8, self)", "5.00", "3.38", "5.00", "5.00"],
            ["Claude → Gemini (n=30)", "4.87 ± 0.51", "4.00 ± 0.64", "5.00 ± 0.00", "4.97 ± 0.18"],
            ["Gemini → Claude (n=30)", "4.60 ± 0.89", "3.33 ± 0.96", "5.00 ± 0.00", "5.00 ± 0.00"],
        ])
    add_para(doc, "")
    add_bullet(doc, "양 LLM 모두 factual ≥ 4.6 / 5, sensitive_leak 5.0 / 5 만점")
    add_bullet(doc, "Gemini self-judge가 자기 비판 방향: comp 3.38 < Claude judge의 4.0")
    add_bullet(doc, "본 시스템의 마스킹 정책이 두 LLM 모두에서 작동")
    add_figure(doc, FIG_DIR / "26_cross_llm_geval.png",
                "그림 9-1. Cross-LLM G-Eval (Judge × Target)")

    add_heading(doc, "9.3 Counterfactual Test 정량화", level=2)
    add_para(doc, "각 인스턴스 30개에서 SHAP top driver 1개를 컨텍스트에서 제거 후 "
                  "LLM 재호출 → 원본과 비교. 변화가 명확할수록 LLM이 컨텍스트에 의존.")
    add_table_from_data(doc,
        headers=["LLM", "Cosine sim (mean ± std)", "ROUGE-L F1"],
        rows=[
            ["Claude Sonnet 4.5", "0.909 ± 0.069", "0.747 ± 0.112"],
            ["Gemini 2.5 Flash", "0.920 ± 0.040", "0.750 ± 0.116"],
        ])
    add_para(doc, "")
    add_para(doc, "두 LLM 모두 cosine 0.91~0.92 — top driver 1개 제거 시 의미는 90% 유지, "
                  "어휘 25%가량 변경. 부분 perturbation에 부분적으로 반응하면서 일관성 유지.")
    add_figure(doc, FIG_DIR / "24_counterfactual_anthropic.png",
                "그림 9-2. Counterfactual cosine + ROUGE-L 분포 (Claude)")

    add_heading(doc, "9.4 Robustness 평가 (프롬프트 변형)", level=2)
    add_para(doc, "각 인스턴스 20개에서 프롬프트 3가지 변형 후 응답 일관성 측정.")
    add_table_from_data(doc,
        headers=["Variant", "Claude cosine", "Gemini cosine"],
        rows=[
            ["role_swap", "0.923 ± 0.057", "0.951 ± 0.034"],
            ["example_swap", "0.914 ± 0.060", "0.908 ± 0.077"],
            ["driver_shuffle", "0.924 ± 0.054", "0.942 ± 0.032"],
        ])
    add_para(doc, "")
    add_para(doc, "두 LLM 모두 6 케이스 전부 cosine ≥ 0.90 (목표 0.85+ 달성). "
                  "본 시스템 운영 안정성 입증.", bold=True)

    add_heading(doc, "9.5 Step 2-A 핵심 메시지", level=2)
    add_bullet(doc, "표본 10배 확장(10→100) 후에도 환각률 0% — 통계적 안정성 입증")
    add_bullet(doc, "Cross-LLM 양방향 평가에서 sensitive_leak 5.0/5 만점 — 마스킹 정책 LLM 종속성 없음")
    add_bullet(doc, "Counterfactual + Robustness 정량화로 시스템 안정성 + 컨텍스트 의존도 측정")
    add_bullet(doc, "Step 1의 핵심 메시지(\"환각 0%, baseline 45.5%\")가 통계적으로 견고함을 입증")

    # ── 10. Step 3-B 보조 테이블 활용 (NEW, 2026-05-05) ──
    add_heading(doc, "10. Step 3-B — 보조 테이블 활용 (성능 확장, 보너스)", level=1)
    add_para(doc, "계획서에 명시된 future work 중 가장 임팩트가 큰 보조 테이블 활용을 D-5에 "
                  "추가 진행. Home Credit Default Risk의 6개 보조 테이블 중 임팩트가 큰 2개"
                  "(bureau, previous_application)를 SK_ID_CURR 단위로 집계해 main 테이블에 "
                  "left merge한 뒤 동일 전처리 정책으로 학습.")

    add_heading(doc, "10.1 데이터 사용 범위와 메모리 효율", level=2)
    add_table_from_data(doc,
        headers=["테이블", "행", "메모리(downcast)", "main 커버리지"],
        rows=[
            ["bureau", "1,716,428", "85 MB", "85.69%"],
            ["bureau_balance", "27,299,925", "156 MB", "(bureau의 45.11%)"],
            ["previous_application", "1,670,214", "147 MB", "94.65%"],
        ])
    add_para(doc, "")
    add_para(doc, "원본 csv 합계 ~950 MB → optimize_dtypes(float32 / int8~32 / category)로 "
                  "388 MB로 -89% 압축. 16 GB RAM 환경에서 5-fold CV 가능.")

    add_heading(doc, "10.2 집계 전략", level=2)
    add_bullet(doc, "bureau_balance → SK_ID_BUREAU (STATUS C/X/0~5 비율 + MONTHS_BALANCE 통계)")
    add_bullet(doc, "bureau (with bb merge) → SK_ID_CURR (전체 / Active / Closed 분리, 346 features)")
    add_bullet(doc, "previous_application → SK_ID_CURR (전체 / Approved / Refused 분리, 410 features)")
    add_bullet(doc, "두 결과 outer merge → 756 aux features")
    add_bullet(doc, "동일 전처리(A1/B1/C1/D/E1/F)로 최종 1161 features")

    add_heading(doc, "10.3 5-fold CV 결과 (test set, mean ± std)", level=2)
    add_table_from_data(doc,
        headers=["Metric", "Baseline (Step 1)", "+ Aux (Step 3-B)", "Δ", "Δ%"],
        rows=[
            ["AUROC", "0.7587 ± 0.0008", "0.7755 ± 0.0011", "+0.0168", "+2.22%"],
            ["AUPRC", "0.2445 ± 0.0011", "0.2646 ± 0.0015", "+0.0201", "+8.21%"],
            ["KS", "0.3846 ± 0.0015", "0.4146 ± 0.0040", "+0.0301", "+7.81%"],
            ["F1", "0.2698 ± 0.0047", "0.2813 ± 0.0096", "+0.0115", "+4.27%"],
        ])
    add_para(doc, "")
    add_bullet(doc, "AUROC +0.0168은 baseline std (0.0008)의 21배 → 통계적으로 명확한 향상")
    add_bullet(doc, "AUPRC + KS의 큰 비율 개선 — positive 8% 불균형에서 분류력 자체 강화")
    add_bullet(doc, "5 fold 모두 0.7743~0.7771 좁은 범위 — 매우 안정적")
    add_figure(doc, FIG_DIR / "27_cv_aux_comparison.png",
                "그림 10-1. Baseline vs Aux 5-fold CV 비교 (test set)")

    add_heading(doc, "10.4 SHAP top 20 변화", level=2)
    add_para(doc, "단일 모델(train+val 학습)로 5,000 test 샘플에 SHAP 계산. baseline top 20과 비교.")
    add_table_from_data(doc,
        headers=["rank", "신규 진입 feature", "mean(|SHAP|)", "의미"],
        rows=[
            ["12", "PREV_NAME_YIELD_GROUP_high_mean", "0.0499", "이전 high-yield 신청 비율"],
            ["13", "PREV_CNT_PAYMENT_std", "0.0486", "이전 결제 횟수 변동성"],
            ["14", "PREV_NAME_YIELD_GROUP_low_action_mean", "0.0450", "이전 low-yield 신청 비율"],
            ["16", "PREV_DAYS_LAST_DUE_1ST_VERSION_max", "0.0440", "이전 만기일 최댓값"],
            ["20", "PREV_NAME_CONTRACT_STATUS_Refused_mean", "0.0360", "이전 거절 비율 ★"],
        ])
    add_para(doc, "")
    add_bullet(doc, "신규 진입 5개 모두 PREV_* (Home Credit 자체 이력) 기반")
    add_bullet(doc, "Bureau (외부 신용기관)는 top 20 진입 못함 — "
                    "main 테이블의 EXT_SOURCE_1/2/3에 외부 신용 정보가 응축돼 있을 가능성 (future work에서 ablation으로 검증)")
    add_bullet(doc, "최강 신규 신호: PREV_NAME_CONTRACT_STATUS_Refused_mean (이전 거절 비율) — 직관과 일치")
    add_figure(doc, FIG_DIR / "29_shap_top20_overlap.png",
                "그림 10-2. baseline vs aux 모델의 SHAP top 20 비교")

    add_heading(doc, "10.5 Step 3-B 핵심 메시지", level=2)
    add_bullet(doc, "보조 테이블 2개로 AUROC 0.7587 → 0.7755 (+2.22%) — 통계적 명확한 향상")
    add_bullet(doc, "AUPRC +8.21%, KS +7.81% — 신용 평가에서 가장 중요한 분류력 강화")
    add_bullet(doc, "외부 신용기관 vs 자체 이력 — 자체 이력의 SHAP 임팩트가 더 큼 (의외 발견)")
    add_bullet(doc, "1161 feature 학습 시간 138s/fold — 메모리·시간 측면에서 운영 가능")

    # ── 11. Step 3-C-1: TabNet 어텐션 × SHAP 융합 컨텍스트 (NEW) ──
    add_heading(doc, "11. Step 3-C-1 — TabNet 어텐션 × SHAP 융합 컨텍스트 (TabNet 통합)",
                  level=1)
    add_para(doc, "기존 4단계 파이프라인의 컨텍스트 빌더 단계에 TabNet의 instance-level "
                  "어텐션 마스크(M_explain)를 통합. SHAP top-k와 Attention top-k를 "
                  "agreed/shap_only/attention_only 3그룹으로 분류해 LLM에 의미 라벨로 "
                  "전달. TabNet이 단순 비교 모델에서 본 메커니즘의 핵심 구성 요소로 통합됨.")

    add_heading(doc, "11.1 Agreement-aware 융합 메커니즘", level=2)
    add_bullet(doc, "agreed_drivers: SHAP top-10과 Attention top-5의 교집합 (두 모델 동의 강한 신호)")
    add_bullet(doc, "shap_only_drivers: SHAP만 본 변수 (부호 + 기여도 보존)")
    add_bullet(doc, "attention_only_drivers: TabNet attention만 본 변수 (sparse, 부호 없음)")
    add_bullet(doc, "LLM 프롬프트에 그룹 라벨 의미 명시 → agreed를 주요 사유로, 보완 그룹은 별도 표현")

    add_heading(doc, "11.2 Agreement 통계 (n=100)", level=2)
    add_table_from_data(doc,
        headers=["그룹", "평균 변수 수", "의미"],
        rows=[
            ["agreed_drivers", "2.12", "두 모델 동의 강한 신호"],
            ["shap_only_drivers", "6.98", "SHAP만 (부호 보존)"],
            ["attention_only_drivers", "2.06", "TabNet만 (부호 없음, sparse)"],
        ])
    add_para(doc, "")
    add_bullet(doc, "n_agreed 분포: 0개=1, 1개=8, 2개=69, 3개=22, 4개+=0 (100 인스턴스)")
    add_bullet(doc, "두 해석 모델이 거의 항상 부분적으로만 겹침 — Day 4의 어텐션-SHAP global ρ=0.117 분석과 일치하는 instance-level 패턴")

    add_heading(doc, "11.3 평가 결과 (n=30 each, judge=Claude)", level=2)
    add_table_from_data(doc,
        headers=["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        rows=[
            ["Halluc strict (↓)", "Anthropic", "0.000", "0.000", "0"],
            ["Halluc strict (↓)", "Gemini", "0.000", "0.000", "0"],
            ["G-Eval Completeness (↑)", "Anthropic", "4.30", "4.97", "+0.67"],
            ["G-Eval Completeness (↑)", "Gemini", "3.90", "4.70", "+0.80"],
            ["G-Eval Factual (↑)", "Anthropic", "4.87", "4.90", "+0.03"],
            ["G-Eval Factual (↑)", "Gemini", "4.90", "4.77", "-0.13"],
            ["G-Eval Sensitive (↑)", "Both", "5.00", "5.00", "0"],
            ["val_match_rate (↑)", "Anthropic", "0.85", "0.90", "+0.06"],
            ["val_match_rate (↑)", "Gemini", "0.79", "0.86", "+0.07"],
        ])
    add_para(doc, "")
    add_bullet(doc, "Halluc 0/30 — 양 LLM × 양 mode 모두 환각 차단 유지 ✅")
    add_bullet(doc, "Completeness +0.67~+0.80 — 두 해석 신호의 상보성이 더 완결한 설명 생성 ★")
    add_bullet(doc, "Factual ≈ 유지, Sensitive 5.0 만점 유지 — 사실성·민감도 손상 없음")
    add_figure(doc, FIG_DIR / "30_fusion_vs_shaponly.png",
                "그림 11-1. SHAP-only vs Fusion 비교 (양 LLM × 4 메트릭)")

    add_heading(doc, "11.4 Step 3-C-1 핵심 메시지", level=2)
    add_bullet(doc, "TabNet이 비교 모델 → 메커니즘 핵심으로 격상 (논문 제목과 본문 일치)")
    add_bullet(doc, "두 해석 신호의 부분 일관 + 부분 상보를 LLM에 명시 노출하는 agreement-aware 컨텍스트 제안")
    add_bullet(doc, "환각 차단 메커니즘 견고 + 완결성 큰 향상의 동시 달성")
    add_bullet(doc, "룰의 sign_match 한계는 다음 절(11.5/Step 3-C-2)에서 NLI로 직접 입증")

    add_heading(doc, "11.5 Step 3-C-2 — NLI 기반 Faithfulness로 평가 객관성 보강",
                  level=2)
    add_para(doc, "Step 3-C-1의 룰 sign_match 하락(Anthropic 0.87→0.65, Gemini 0.94→0.77)이 "
                  "키워드 한계 때문임을 의미적 측정으로 입증하기 위해, 다국어 NLI 모델"
                  "(KLUE 학습 데이터 100M+ 포함 mDeBERTa-v3-xnli)을 사용해 자동 의미 함의 "
                  "평가를 추가했다.")
    add_para(doc, "원리: 컨텍스트 facts를 자연어 premise로 변환, LLM 설명을 문장 단위로 split "
                  "후 각 (premise, hypothesis) 쌍에 대해 entailment / neutral / contradiction "
                  "확률 계산. 인스턴스별 entailment_rate 평균.")
    add_table_from_data(doc,
        headers=["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        rows=[
            ["entailment_rate (↑)", "Anthropic", "0.413", "0.625", "+0.212 ★"],
            ["entailment_rate (↑)", "Gemini", "0.509", "0.624", "+0.115 ★"],
            ["contradiction_rate (↓)", "Anthropic", "0.366", "0.191", "-0.175 ★"],
            ["contradiction_rate (↓)", "Gemini", "0.307", "0.167", "-0.140 ★"],
            ["min_entailment (↑)", "Anthropic", "0.048", "0.181", "+0.134"],
            ["min_entailment (↑)", "Gemini", "0.086", "0.082", "-0.004 ≈"],
        ])
    add_para(doc, "")
    add_bullet(doc, "양 LLM 모두 entailment 향상 + contradiction 감소 — fusion이 의미적으로 더 충실함을 확정")
    add_bullet(doc, "룰의 sign_match 하락은 키워드 셋의 한계임이 NLI로 직접 입증 — fusion에서 LLM이 \"증가시키는\", \"위험\" 등 다양한 표현 사용해도 의미적 함의는 정확히 유지")
    add_bullet(doc, "3-tier 평가 체계 완성: Rules(키워드) + G-Eval(LLM judge) + NLI(의미 함의). 본 논문 약점 1번(LLM 평가 객관성) 부분 해소")
    add_figure(doc, FIG_DIR / "31_nli_vs_rules.png",
                "그림 11-2. NLI Entailment / Contradiction / Rule sign_match 3-패널 비교")

    add_heading(doc, "11.6 Step 3-C-2-f — Cross-Judge G-Eval 검증 (보너스)",
                  level=2)
    add_para(doc, "Step 3-C-1/3-C-2의 G-Eval 결과는 Claude judge 단독으로 측정. "
                  "Step 2-A 패턴의 양방향 cross-LLM judge를 fusion 평가에도 적용해 "
                  "judge 종속성을 검증.")
    add_para(doc, "동일 4 그룹(SHAP-only/Fusion × Anthropic/Gemini target, n=30 each)을 "
                  "Gemini judge로 재평가. Gemini API 503 과부하로 retry 로직 강화 "
                  "(30s/60s/120s/240s 백오프). 총 transient retry 40회, fail 0회, 모든 "
                  "120 호출 회복.")

    add_table_from_data(doc,
        headers=["Target", "Metric", "Δ Claude judge", "Δ Gemini judge", "차이"],
        rows=[
            ["Anthropic", "Completeness", "+0.667", "+0.900", "0.23"],
            ["Anthropic", "Factual", "+0.033", "+0.133", "0.10"],
            ["Anthropic", "Sensitive", "0", "0", "0"],
            ["Gemini", "Completeness", "+0.800", "+1.100", "0.30"],
            ["Gemini", "Factual", "-0.133", "+0.467", "0.60 ★"],
            ["Gemini", "Sensitive", "0", "0", "0"],
        ])
    add_para(doc, "")
    add_bullet(doc, "Completeness: 양 judge에서 모두 큰 향상 (Claude +0.67/+0.80, Gemini +0.90/+1.10) — fusion 효과가 judge 종속 아님 명확히 입증")
    add_bullet(doc, "Gemini target Factual: Claude judge -0.13 (사실상 동등) vs Gemini judge +0.47 (명확한 향상). 차이 0.60 — **단일 judge였으면 잘못된 결론 위험**, Cross-LLM judge의 가치 직접 입증")
    add_bullet(doc, "Sensitive Leak: 양 judge × 양 mode 모두 5.0/5.0 만점 — 마스킹 정책의 LLM 종속성 없음 다층 검증")
    add_bullet(doc, "Self-bias 패턴 약함 — Gemini judge가 self/cross 모두에 더 큰 Δ를 부여하는 패턴")
    add_figure(doc, FIG_DIR / "32_cross_judge_geval.png",
                "그림 11-3. Cross-Judge G-Eval (Claude + Gemini judge × SHAP-only/Fusion)")

    # ── 12. Future Work ──
    add_heading(doc, "12. 향후 계획 (Future Work)", level=1)
    add_bullet(doc, "인간 평가 (Plausibility) — IRB 간소판, 5점 리커트 척도, Cohen's κ 신뢰도 측정 ★ 1순위")
    add_bullet(doc, "Fusion 표본 30 → 100 확장 + counterfactual 결합")
    add_bullet(doc, "3-way ablation: SHAP-only vs Attention-only vs Fusion")
    add_bullet(doc, "잔여 보조 테이블 4개(POS_CASH_balance, credit_card_balance, installments_payments, bureau_balance 추가 활용) → AUROC 0.78+ 추가 향상")
    add_bullet(doc, "Bureau ablation — bureau 단독 vs prev 단독 vs 둘 다 비교로 EXT_SOURCE 응축 가설 검증")
    add_bullet(doc, "TabNet, LightGBM에 aux 효과 일반화 (Step 3-B는 XGBoost만 검증)")
    add_bullet(doc, "본격 fairness-aware 학습 — Reweighing, Adversarial Debiasing")
    add_bullet(doc, "FT-Transformer 비교 모델 추가")
    add_bullet(doc, "한국어 native NLI 모델 추가 검증 (현재는 다국어 NLI; torch 환경 정비 후 KLUE-roberta-NLI)")
    add_bullet(doc, "한국어 도메인 특화 금융 LLM 미세조정 (QLoRA)")
    add_bullet(doc, "(완료) Gemini judge로 양방향 G-Eval cross-validation — Step 3-C-2-f에서 진행, 양 judge 일관 입증")

    # 저장
    out = PAPER_DIR / "midterm_report.docx"
    doc.save(out)
    print(f"[OK] {out} 저장")


if __name__ == "__main__":
    main()
