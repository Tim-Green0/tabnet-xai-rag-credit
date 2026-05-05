"""연구 결과 정리 docx 생성 (일회성, 날짜 prefix).

Step 1 ~ Step 3-C-2-f의 전체 흐름과 결과를 처음부터 깔끔하게 한 문서로 정리.
기존 midterm_report*.docx와 별개로, 한 번에 보는 종합 정리용.

산출:
    paper/{YYYY-MM-DD}_연구결과정리.docx

사용:
    PYTHONIOENCODING=utf-8 .venv/Scripts/python.exe -m src.gen_summary_docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from src.gen_friendly_report import (
    H, P, Quote, Bullet, Table, Fig, PageBreak,
)
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PAPER_DIR = Path("paper")
FIG_DIR = Path("figures")


def main() -> None:
    today = date.today().isoformat()  # 2026-05-05
    out = PAPER_DIR / f"{today}_연구결과정리.docx"

    doc = Document()
    # 본문 폰트 기본
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    # ─────────────────────────────────────
    # 표지
    # ─────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TabNet × SHAP × LLM 기반 XAI-RAG 신용 평가")
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("— 환각 없는 자연어 설명을 위한 두 해석 신호 융합 RAG —")
    r2.font.name = "Malgun Gothic"
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    P(doc, "")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run(
        f"연구 결과 정리 (Step 1 ~ Step 3-C-2-f)  |  작성일: {today}\n"
        "전공: 데이터사이언스 · 인공지능 석사  |  학번 A70067 · 오현택  |  지도교수: 박운상"
    )
    r3.font.name = "Malgun Gothic"
    r3.font.size = Pt(11)

    P(doc, "")
    Quote(doc, "본 문서는 미팅용 자료가 아니라, 현재 시점(2026-05-05)까지 진행한 "
                "연구의 전체 흐름과 결과를 한 번에 보기 위한 종합 정리. "
                "기존 paper/midterm_report*.docx, midterm_slides.pptx와 별개의 일회성 자료.")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 1. 연구 개요
    # ─────────────────────────────────────
    H(doc, "1. 연구 개요", level=1)
    P(doc, "본 연구는 정형 데이터(신용 평가)에서 다음 세 가지를 동시에 달성하는 통합 "
            "파이프라인을 제안한다:")
    Bullet(doc, "예측 성능 — TabNet과 트리 모델(XGBoost/LightGBM)로 SOTA에 준하는 분류 성능")
    Bullet(doc, "투명한 근거 — SHAP과 TabNet 어텐션의 상보성을 정량 분석")
    Bullet(doc, "신뢰할 수 있는 자연어 설명 — LLM이 SHAP/Attention의 검색된 근거(retrieved evidence)만 받아쓰는 XAI-RAG로 환각 차단")

    H(doc, "1.1 핵심 아이디어 (한 줄)", level=2)
    Quote(doc, "\"SHAP과 TabNet 어텐션의 결합 결과를 LLM의 retrieved evidence로 재정의하면, "
                "LLM 종속성 없이 환각이 0%인 자연어 설명 리포트를 생성할 수 있다.\"")

    H(doc, "1.2 4단계 파이프라인", level=2)
    P(doc, "[정형 데이터] → [XGBoost 예측 + SHAP local] + [TabNet 어텐션] "
            "→ [Agreement-aware JSON 컨텍스트 (민감변수 마스킹)] "
            "→ [LLM 자연어 설명 생성] → [정량 평가: Rules + G-Eval(Cross-Judge) + NLI]",
       size=10)

    H(doc, "1.3 데이터와 환경", level=2)
    Bullet(doc, "메인 데이터: Kaggle Home Credit Default Risk (application_train.csv 307,511 × 122)")
    Bullet(doc, "보조 데이터: bureau (1.7M), bureau_balance (27.3M), previous_application (1.7M) — Step 3-B에서 추가")
    Bullet(doc, "분할: 60/20/20 stratified, SEED=42")
    Bullet(doc, "환경: Windows 10 / Python 3.10 / GTX 1660 Ti (6GB) / CUDA 12.1")
    Bullet(doc, "LLM API: Gemini 2.5 Flash + Claude Sonnet 4.5")
    Bullet(doc, "NLI 모델: mDeBERTa-v3-base-xnli (다국어, 한국어 학습 데이터 100M+ 포함)")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 2. Step 1 — 미팅 프로토타입
    # ─────────────────────────────────────
    H(doc, "2. Step 1 — 미팅용 작동 프로토타입", level=1)
    P(doc, "8일 작업으로 4단계 파이프라인을 처음부터 끝까지 작동하는 형태로 구축. "
            "표본 10명 기준 정량 평가까지 완료.")

    H(doc, "2.1 모델 비교 (5-fold CV, test set)", level=2)
    Table(doc,
        ["모델", "AUROC (mean ± std)", "비고"],
        [["XGBoost (★ 1등)", "0.7587 ± 0.0008", "5/5 fold 최고"],
         ["LightGBM", "0.7549 ± —", ""],
         ["Logistic Regression", "0.7547 ± —", ""],
         ["TabNet (tuned)", "0.7543 ± —", "어텐션 해석성"]])

    H(doc, "2.2 SHAP × TabNet 어텐션 일관성 (RQ2)", level=2)
    Bullet(doc, "Spearman ρ (전체 214 features) = 0.117 — 약한 양의 상관")
    Bullet(doc, "Top-50 union ρ = -0.195 — 미세 구조에서 상보적")
    Bullet(doc, "Top-20 Jaccard = 0.29 — 핵심 변수는 일관, 미세는 다름")
    Quote(doc, "→ 두 해석 신호가 \"부분 일관 + 부분 상보\"라는 결론. "
                "이 특성은 Step 3-C-1의 융합 컨텍스트 동기가 됨.")

    H(doc, "2.3 공정성 진단 (Day 5)", level=2)
    Bullet(doc, "4 모델 × {GENDER, AGE} = 8/8 케이스에서 4/5 rule 위반 (Disparate Impact < 0.8)")
    Bullet(doc, "GENDER ablation은 효과적 (DP -36~40%, AUROC -0.005~-0.010)")
    Bullet(doc, "AGE ablation은 효과 미미 — proxy variable로 간접 인코딩됨 (DAYS_EMPLOYED 등)")

    H(doc, "2.4 XAI-RAG 자연어 설명 + 정량 평가 (RQ3)", level=2)
    Bullet(doc, "Demo idx 59291 (True Positive): 정형 데이터 → XGBoost (P=0.948) → SHAP top 5 → JSON 컨텍스트 → Gemini/Claude 자연어 설명")
    Bullet(doc, "Hallucination Rate (10 샘플) = 0.000 (양 LLM 모두)")
    Bullet(doc, "Faithfulness: 변수명·값·SHAP 부호 정확 인용")
    Bullet(doc, "민감변수 마스킹 (CODE_GENDER, DAYS_BIRTH 등) 작동 — 직접 인용 0건")

    H(doc, "2.5 ★ Counterfactual Baseline — Step 1의 결정타", level=2)
    P(doc, "\"SHAP 컨텍스트가 정말 환각을 막아주는가?\"를 직접 검증.")
    Bullet(doc, "XAI-RAG (with SHAP context): Hallucination 0%")
    Bullet(doc, "no-SHAP baseline (raw 데이터만): Claude 45.5% — DTI/LTV/DSR/햇살론·미소금융/가짜 전화번호 1588-XXXX 등 환각")
    Bullet(doc, "Gemini는 raw에서도 변수명만 인용 → 측정상 0%지만 \"환각이 없는 게 아니라 잡히지 않은 것\"")
    Quote(doc, "SHAP 컨텍스트의 유무가 환각률에 결정적 차이를 만든다는 직접 증거. "
                "Step 1의 가장 강력한 메시지.")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 3. Step 2-A — 평가 신뢰성
    # ─────────────────────────────────────
    H(doc, "3. Step 2-A — 평가 신뢰성 강화", level=1)
    P(doc, "Step 1의 가장 큰 약점이었던 \"표본 10명\"을 100명으로 확장하고, Cross-LLM G-Eval, "
            "Counterfactual 정량화, Robustness 평가를 추가.")

    H(doc, "3.1 100명 표본 환각률", level=2)
    Bullet(doc, "Gemini 2.5 Flash: Halluc 0/100")
    Bullet(doc, "Claude Sonnet 4.5: Halluc 0/100")
    Bullet(doc, "→ Step 1의 10명 결과가 통계적으로 견고함을 입증")

    H(doc, "3.2 Cross-LLM G-Eval (양방향, n=30 each)", level=2)
    Table(doc,
        ["Judge → Target", "Factual", "Completeness", "Sensitive", "Style"],
        [["Claude → Gemini", "4.87 ± 0.51", "4.00 ± 0.64", "5.00 ± 0.00", "4.97 ± 0.18"],
         ["Gemini → Claude", "4.60 ± 0.89", "3.33 ± 0.96", "5.00 ± 0.00", "5.00 ± 0.00"]])
    Bullet(doc, "양 LLM 모두 factual ≥ 4.6, sensitive 5.0 만점 → 마스킹 정책 LLM 무관 견고")

    H(doc, "3.3 Counterfactual + Robustness 정량화", level=2)
    Bullet(doc, "Counterfactual (top driver 1개 제거, n=30): cosine 0.909~0.920, ROUGE-L 0.747~0.750 — 부분 perturbation에 부분 반응 + 의미 일관 유지")
    Bullet(doc, "Robustness (3 변형 × n=20): cosine 0.908~0.951 — 프롬프트 미세 변형에 강건")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 4. Step 3-B — 성능 확장 (보조 테이블)
    # ─────────────────────────────────────
    H(doc, "4. Step 3-B — 성능 확장 (보조 테이블 활용)", level=1)
    P(doc, "Home Credit 보조 테이블 6개 중 임팩트가 큰 2개 (bureau, previous_application)를 "
            "SK_ID_CURR 단위로 집계해 main에 left merge. 동일 전처리 정책으로 학습.")

    H(doc, "4.1 메모리 효율 + 집계", level=2)
    Bullet(doc, "원본 csv 합계 ~950MB → optimize_dtypes로 388MB (-89%)")
    Bullet(doc, "bureau (with bureau_balance merge): 346 features (전체/Active/Closed 분리)")
    Bullet(doc, "previous_application: 410 features (전체/Approved/Refused 분리)")
    Bullet(doc, "Outer merge → 756 aux features → 동일 전처리로 최종 1161 features")

    H(doc, "4.2 5-fold CV 결과 (XGBoost, test set)", level=2)
    Table(doc,
        ["Metric", "Baseline (Step 1)", "+ Aux (Step 3-B)", "Δ"],
        [["AUROC", "0.7587 ± 0.0008", "0.7755 ± 0.0011", "+0.0168 (+2.22%)"],
         ["AUPRC", "0.2445 ± 0.0011", "0.2646 ± 0.0015", "+0.0201 (+8.21%)"],
         ["KS", "0.3846 ± 0.0015", "0.4146 ± 0.0040", "+0.0301 (+7.81%)"],
         ["F1", "0.2698 ± 0.0047", "0.2813 ± 0.0096", "+0.0115 (+4.27%)"]])
    Bullet(doc, "AUROC +0.0168은 baseline std (0.0008)의 21배 → 통계적 명확한 향상")
    Bullet(doc, "AUPRC + KS의 큰 비율 개선 — positive 8% 불균형에서 분류력 자체 강화")
    Fig(doc, FIG_DIR / "27_cv_aux_comparison.png",
        "그림 4-1. Baseline vs Aux 5-fold CV 비교 (test set)")

    H(doc, "4.3 SHAP top 20 변화 — 의외의 발견", level=2)
    Bullet(doc, "신규 진입 5개 모두 PREV_* (Home Credit 자체 신청 이력) 기반")
    Bullet(doc, "최강 신규 신호: PREV_NAME_CONTRACT_STATUS_Refused_mean — 이전 거절 비율, 직관과 일치")
    Bullet(doc, "★ Bureau (외부 신용기관) feature는 top 20 진입 못함 — main의 EXT_SOURCE_1/2/3에 응축돼 있을 가능성 (future work에서 ablation으로 검증)")
    Fig(doc, FIG_DIR / "29_shap_top20_overlap.png",
        "그림 4-2. baseline vs aux 모델의 SHAP top 20 비교")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 5. Step 3-C-1 — TabNet 통합
    # ─────────────────────────────────────
    H(doc, "5. Step 3-C-1 — TabNet 어텐션 × SHAP 융합 컨텍스트", level=1)
    P(doc, "Step 1의 큰 약점 — 논문 제목에 \"TabNet\"이 들어가지만 실제로는 비교 모델 "
            "역할만 했다는 점 — 을 정면 대응. TabNet의 instance-level 어텐션을 SHAP과 융합한 "
            "agreement-aware 컨텍스트를 LLM에 제공하는 메커니즘 구현.")

    H(doc, "5.1 융합 메커니즘 — 3 그룹 라벨", level=2)
    Bullet(doc, "agreed_drivers: SHAP top-10 ∩ Attention top-5 (두 모델 동의 강한 신호)")
    Bullet(doc, "shap_only_drivers: SHAP만 본 변수 (부호 + 기여도 보존)")
    Bullet(doc, "attention_only_drivers: TabNet만 본 변수 (sparse, 부호 없음)")
    Bullet(doc, "LLM 프롬프트에 그룹 라벨 의미 명시 → \"동의\" 신호를 주요 사유로, \"보완\" 신호를 부가 표현")

    H(doc, "5.2 Agreement 통계 (n=100)", level=2)
    Bullet(doc, "평균 agreed=2.12, shap_only=6.98, attention_only=2.06")
    Bullet(doc, "n_agreed 분포: 0개=1, 1개=8, 2개=69, 3개=22, 4개+=0")
    Quote(doc, "두 해석 모델은 거의 항상 부분적으로만 겹친다 (3+ 동의 22%, 4+ 동의 0%) — "
                "Day 4의 ρ=0.117 분석과 일치하는 instance-level 패턴.")

    H(doc, "5.3 결과 (n=30 each, Claude judge)", level=2)
    Table(doc,
        ["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        [["Halluc strict (↓)", "Anthropic", "0.000", "0.000", "0 ✅"],
         ["Halluc strict (↓)", "Gemini", "0.000", "0.000", "0 ✅"],
         ["G-Eval Completeness (↑)", "Anthropic", "4.30", "4.97", "+0.67 ★"],
         ["G-Eval Completeness (↑)", "Gemini", "3.90", "4.70", "+0.80 ★"],
         ["G-Eval Factual (↑)", "Anthropic", "4.87", "4.90", "+0.03 ≈"],
         ["G-Eval Factual (↑)", "Gemini", "4.90", "4.77", "-0.13 ≈"],
         ["Sensitive Leak (↑)", "Both", "5.00", "5.00", "0 ✅"]])
    Bullet(doc, "환각 차단 메커니즘 견고 + 완결성 큰 향상의 동시 달성")
    Bullet(doc, "TabNet이 단순 비교 모델 → 메커니즘 핵심으로 격상 (논문 제목 본문 일치)")
    Fig(doc, FIG_DIR / "30_fusion_vs_shaponly.png",
        "그림 5-1. SHAP-only vs Fusion 비교 (양 LLM × 4 메트릭)")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 6. Step 3-C-2 — NLI 평가
    # ─────────────────────────────────────
    H(doc, "6. Step 3-C-2 — NLI 기반 Faithfulness 보강", level=1)
    P(doc, "Step 3-C-1의 룰 sign_match 하락(Anthropic 0.87→0.65, Gemini 0.94→0.77)이 "
            "키워드 한계 때문임을 의미적 측정으로 입증하기 위해, 다국어 NLI 모델을 사용해 "
            "자동 함의 평가 추가.")

    H(doc, "6.1 NLI 평가 알고리즘", level=2)
    Bullet(doc, "Premise = LLM에 주어진 컨텍스트 facts를 자연어 단락으로 변환")
    Bullet(doc, "Hypothesis = LLM 설명의 각 문장 (advice/disclaimer 섹션 제외)")
    Bullet(doc, "NLI 모델이 (premise, hypothesis) 쌍에 대해 entailment/neutral/contradiction 확률 출력")
    Bullet(doc, "인스턴스별 entailment_rate 평균 → faithfulness score")

    H(doc, "6.2 결과 (n=30 each)", level=2)
    Table(doc,
        ["Metric", "LLM", "SHAP-only", "Fusion", "Δ"],
        [["entailment_rate (↑)", "Anthropic", "0.413", "0.625", "+0.212 ★"],
         ["entailment_rate (↑)", "Gemini", "0.509", "0.624", "+0.115 ★"],
         ["contradiction_rate (↓)", "Anthropic", "0.366", "0.191", "-0.175 ★"],
         ["contradiction_rate (↓)", "Gemini", "0.307", "0.167", "-0.140 ★"],
         ["min_entailment (↑)", "Anthropic", "0.048", "0.181", "+0.134"]])
    Bullet(doc, "양 LLM 모두 entailment 향상 + contradiction 감소 — fusion이 의미적으로 더 충실")
    Bullet(doc, "★ 룰 sign_match 하락은 키워드 셋의 한계임이 NLI로 직접 입증 (다양한 표현 사용해도 의미 함의는 정확히 유지)")
    Bullet(doc, "3-tier 평가 체계 완성: Rules + G-Eval + NLI")
    Fig(doc, FIG_DIR / "31_nli_vs_rules.png",
        "그림 6-1. NLI Entailment / Contradiction / Rule sign_match 3-패널 비교")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 7. Step 3-C-2-f — Cross-Judge
    # ─────────────────────────────────────
    H(doc, "7. Step 3-C-2-f — Cross-Judge G-Eval 검증", level=1)
    P(doc, "Step 3-C-2까지의 G-Eval은 Claude judge 단독. Step 2-A의 양방향 cross-LLM judge "
            "패턴을 fusion 평가에도 적용해 judge 종속성 검증.")

    H(doc, "7.1 결과 (Δ = fusion - shaponly, n=30 each)", level=2)
    Table(doc,
        ["Target", "Metric", "Δ Claude judge", "Δ Gemini judge", "차이"],
        [["Anthropic", "Completeness", "+0.667 ★", "+0.900 ★", "0.23"],
         ["Anthropic", "Factual", "+0.033", "+0.133", "0.10"],
         ["Anthropic", "Sensitive", "0", "0", "0"],
         ["Gemini", "Completeness", "+0.800 ★", "+1.100 ★", "0.30"],
         ["Gemini", "Factual", "-0.133", "+0.467", "0.60 ★★"],
         ["Gemini", "Sensitive", "0", "0", "0"]])

    H(doc, "7.2 핵심 발견", level=2)
    Bullet(doc, "Completeness: 양 judge 모두 큰 향상 (Claude +0.67/+0.80, Gemini +0.90/+1.10) — fusion 효과 judge 종속 아님")
    Bullet(doc, "★ Gemini target Factual: Claude judge -0.13 vs Gemini judge +0.47, 차이 0.60 — 단일 judge였으면 잘못된 결론 위험. Cross-LLM judge의 가치 직접 입증")
    Bullet(doc, "Sensitive Leak: 양 judge × 양 mode 모두 5.0/5.0 만점 — 마스킹 정책 견고 다층 검증")
    Bullet(doc, "Self-bias 패턴 약함 — Gemini judge가 self/cross 모두에 더 큰 Δ 부여")
    Fig(doc, FIG_DIR / "32_cross_judge_geval.png",
        "그림 7-1. Cross-Judge G-Eval (4 메트릭 × 2 target × 2 judge × 2 mode)")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 8. 종합 메시지
    # ─────────────────────────────────────
    H(doc, "8. 종합 — 본 연구의 5가지 가치", level=1)

    H(doc, "8.1 환각 차단 메커니즘", level=2)
    Bullet(doc, "Halluc 0/10 (Step 1) → 0/100 (Step 2-A) → 0/30 (Step 3-C-1 fusion) — 통계적 견고")
    Bullet(doc, "Counterfactual baseline: Claude 45.5% (no-SHAP) vs 0% (XAI-RAG) — 직접 증거")
    Bullet(doc, "NLI contradiction -0.14~-0.18 (fusion 추가) — 의미적 측정에서도 견고")

    H(doc, "8.2 예측 성능 (Step 1 → 3-B)", level=2)
    Bullet(doc, "XGBoost AUROC 0.7587 → 0.7755 (+2.22%)")
    Bullet(doc, "AUPRC +8.21%, KS +7.81% — 불균형 분류력 강화")

    H(doc, "8.3 해석 융합", level=2)
    Bullet(doc, "어텐션 × SHAP 부분 일관 + 부분 상보 (ρ=0.117 global, instance-level n_agreed 평균 2.12)")
    Bullet(doc, "agreement-aware 컨텍스트로 LLM에 의미 라벨 전달")
    Bullet(doc, "TabNet이 비교 모델 → 메커니즘 핵심으로 격상 (논문 제목 정당화)")

    H(doc, "8.4 완결성·충실성 향상", level=2)
    Bullet(doc, "G-Eval Completeness: Claude judge +0.67/+0.80, Gemini judge +0.90/+1.10")
    Bullet(doc, "NLI Entailment: Anthropic +0.21, Gemini +0.12 — 다층 평가 일관 향상")
    Bullet(doc, "Factual ≈ 유지 (Cross-judge로 정밀 검증), Sensitive 5.0/5.0 양 judge 만점")

    H(doc, "8.5 평가 신뢰성", level=2)
    Bullet(doc, "표본: 10 → 100 (Halluc) / 30 (Cross-LLM G-Eval, Counterfactual, Robustness, NLI)")
    Bullet(doc, "3-tier 평가: Rules (룰) + G-Eval (LLM-as-judge × Cross-LLM) + NLI (의미 함의)")
    Bullet(doc, "Robustness: 프롬프트 미세 변형에 cosine 0.91~0.95 강건")

    P(doc, "")
    P(doc, "★ 한 줄 메시지", bold=True, size=14)
    P(doc, "\"본 XAI-RAG 시스템은 두 해석 신호(SHAP + TabNet 어텐션)의 상보성을 LLM에게 "
            "명시 라벨로 제공함으로써, 사실성·민감변수·완결성을 손상시키지 않으면서 환각 0%를 "
            "양 LLM에서 유지하며, 룰·G-Eval(Cross-Judge)·NLI 3-tier 평가에서 일관된 향상을 "
            "보인다.\"", bold=True, size=12)

    PageBreak(doc)

    # ─────────────────────────────────────
    # 9. 한계와 향후 계획
    # ─────────────────────────────────────
    H(doc, "9. 한계와 향후 계획", level=1)

    H(doc, "9.1 현재 한계 (정직한 인식)", level=2)
    Bullet(doc, "Fusion 평가 표본 30명 — 100명+ 확장 검토")
    Bullet(doc, "인간평가 (Plausibility) 미수행 — 자동 평가만으로는 \"사람이 정말 신뢰하는가\" 검증 부족")
    Bullet(doc, "보조 테이블 2/6개 사용 (Step 3-B) — 4개 잔여")
    Bullet(doc, "Bureau의 SHAP 미진입 — EXT_SOURCE 응축 가설 ablation으로 검증 필요")
    Bullet(doc, "TabNet-only 컨텍스트 ablation 미수행 (3-way 비교 미완)")
    Bullet(doc, "TabNet/LightGBM에 보조 테이블 효과 일반화 미실시")
    Bullet(doc, "Fairness-aware 학습 미수행 (8/8 케이스 4/5 rule 위반 진단만)")
    Bullet(doc, "Generic RAG (도메인 일반 지식) baseline 미비교 — Counterfactual baseline 정당성 보강 필요")
    Bullet(doc, "데이터 단일 (Home Credit) — UCI German Credit 등 다른 데이터셋 미검증")
    Bullet(doc, "한국어 native NLI 모델 추가 검증 필요 (현재 다국어)")

    H(doc, "9.2 향후 계획 우선순위", level=2)
    Bullet(doc, "1순위: 인간평가 (Plausibility) — IRB 간소판, 5점 척도, Cohen's κ — 약점 1번 완전 해소")
    Bullet(doc, "2순위: Fairness-aware 학습 — Reweighing, Adversarial Debiasing — 4/5 rule 통과 시도")
    Bullet(doc, "3순위: Generic RAG baseline — Counterfactual baseline 정당성 보강")
    Bullet(doc, "4순위: 잔여 보조 테이블 4개로 AUROC 0.78+ 추가 향상 + Bureau ablation")
    Bullet(doc, "5순위: 데이터 다양성 (UCI German Credit) — 일반화 입증")
    Bullet(doc, "6순위: 3-way ablation, TabNet/LightGBM 일반화, FT-Transformer 비교")
    Bullet(doc, "장기: 한국어 도메인 특화 LLM 미세조정 (QLoRA)")

    PageBreak(doc)

    # ─────────────────────────────────────
    # 10. 산출물 인덱스
    # ─────────────────────────────────────
    H(doc, "10. 산출물 인덱스 (참고)", level=1)

    H(doc, "10.1 git tag 흐름", level=2)
    Bullet(doc, "step1 (eaef964) — Day 8, 미팅 프로토타입 8일차")
    Bullet(doc, "step2a (5cd2192) — 100명, Cross-LLM, Robustness")
    Bullet(doc, "step3 (= step3c1, 1fe2e24) — 보조 테이블 + TabNet 융합")
    Bullet(doc, "step3b (7427c39) — 보조 테이블 단독")
    Bullet(doc, "step3c2 (4c48ab0) — NLI")
    Bullet(doc, "step4 (5758116) — NLI + Cross-Judge + 미팅 자료 통합 ★ 현재 마일스톤")

    H(doc, "10.2 핵심 figure", level=2)
    Bullet(doc, "figures/13_cv_comparison.png — Step 1 5-fold CV")
    Bullet(doc, "figures/16_attention_vs_shap_scatter.png — 어텐션 vs SHAP")
    Bullet(doc, "figures/22_demo_walkthrough.png — Demo idx 59291")
    Bullet(doc, "figures/27_cv_aux_comparison.png — Step 3-B 보조 테이블 효과")
    Bullet(doc, "figures/29_shap_top20_overlap.png — SHAP top 20 변화")
    Bullet(doc, "figures/30_fusion_vs_shaponly.png — Step 3-C-1 fusion 비교")
    Bullet(doc, "figures/31_nli_vs_rules.png — Step 3-C-2 NLI 평가")
    Bullet(doc, "figures/32_cross_judge_geval.png — Step 3-C-2-f cross-judge")

    H(doc, "10.3 day별 보고서", level=2)
    Bullet(doc, "results/day1~8_summary.md — Step 1 (환경, EDA, 베이스라인, TabNet, SHAP, 공정성, XAI-RAG, 평가, Demo)")
    Bullet(doc, "results/step2a_summary.md — Step 2-A 종합")
    Bullet(doc, "results/day9_summary.md — Step 3-B 보조 테이블")
    Bullet(doc, "results/day10_summary.md — Step 3-C-1 융합")
    Bullet(doc, "results/day11_summary.md — Step 3-C-2 NLI")
    Bullet(doc, "results/day12_summary.md — Step 3-C-2-f Cross-Judge")

    H(doc, "10.4 코드 모듈", level=2)
    Bullet(doc, "src/utils.py, data_loader.py, eda.py, preprocess.py — 데이터/전처리")
    Bullet(doc, "src/baselines.py, cv_eval.py, tabnet_train.py — 모델")
    Bullet(doc, "src/shap_analysis.py, fairness.py — 해석/공정성")
    Bullet(doc, "src/context_builder.py, llm_explainer.py, eval_explanation.py — Step 1 XAI-RAG")
    Bullet(doc, "src/expand_samples.py, cross_llm_geval.py, counterfactual_test.py, robustness_test.py, text_similarity.py — Step 2-A")
    Bullet(doc, "src/aux_data.py, aux_features.py, preprocess_with_aux.py, cv_eval_aux.py, shap_aux.py — Step 3-B")
    Bullet(doc, "src/tabnet_attention_local.py, fusion_context.py, llm_explainer_fusion.py, eval_fusion.py — Step 3-C-1")
    Bullet(doc, "src/nli_eval.py, cross_judge_analysis.py — Step 3-C-2/3-C-2-f")

    P(doc, "")
    Quote(doc, f"본 정리는 {today} 시점의 스냅샷. git에서 step1~step4 tag로 각 시점 재현 가능.")

    # 저장
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[OK] {out} 저장")


if __name__ == "__main__":
    main()
